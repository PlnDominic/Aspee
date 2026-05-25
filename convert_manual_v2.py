"""Convert ASPEE_PHARMA_USER_MANUAL.md  →  ASPEE_PHARMA_USER_MANUAL.docx (v2.0)"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours ──────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1e, 0x3a, 0x5f)
TEAL  = RGBColor(0x0f, 0x76, 0x6e)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY = RGBColor(0xF1, 0xF5, 0xF9)
DGRAY = RGBColor(0x33, 0x33, 0x33)
AMBER = RGBColor(0x92, 0x40, 0x0E)
GREEN = RGBColor(0x16, 0x4e, 0x32)

# ── Helpers ───────────────────────────────────────────────────────────────────

def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '12')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1e3a5f')
    pBdr.append(bot)
    pPr.append(pBdr)

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = TEAL

def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DGRAY

def add_body(doc, text):
    """Body paragraph with inline **bold** support."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = DGRAY
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.color.rgb = DGRAY

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10)

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FEF3C7')
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run('  Note: ' + text)
    run.font.size  = Pt(9.5)
    run.font.color.rgb = AMBER
    run.italic = True

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F5F9')
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = DGRAY

def add_table(doc, headers, rows):
    n = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    hr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        cell_bg(c, '1e3a5f')
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]
        pp.paragraph_format.space_before = Pt(3)
        pp.paragraph_format.space_after  = Pt(3)
        run = pp.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
    # Rows
    for ri, row_data in enumerate(rows):
        r = tbl.rows[ri + 1]
        bg = 'F1F5F9' if ri % 2 == 0 else 'FFFFFF'
        for ci, txt in enumerate(row_data):
            c = r.cells[ci]
            cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            pp = c.paragraphs[0]
            pp.paragraph_format.space_before = Pt(2)
            pp.paragraph_format.space_after  = Pt(2)
            run = pp.add_run(str(txt))
            run.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── Markdown parser ───────────────────────────────────────────────────────────

def parse_inline(text):
    """Strip markdown inline markers for plain text (used in tables)."""
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text.strip()

def parse_table(lines):
    """Parse a markdown table block into headers + rows."""
    headers = []
    rows = []
    for line in lines:
        if re.match(r'^\s*\|[-| :]+\|\s*$', line):
            continue
        cells = [parse_inline(c) for c in line.strip().strip('|').split('|')]
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows

def convert(md_path, doc):
    with open(md_path, encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code = False
    code_lines = []
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # Code fences
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                if code_lines:
                    add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if re.match(r'^\s*\|', line):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                h, r = parse_table(table_lines)
                if h:
                    add_table(doc, h, r)
                table_lines = []

        # Skip YAML-style hr lines
        if re.match(r'^---+\s*$', line):
            i += 1
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            # Skip document title and version lines at top
            if 'ASPEE PHARMACEUTICALS' in text or 'Staff User Manual' in text or 'Version' in text:
                i += 1
                continue
            add_heading1(doc, text)
            i += 1
            continue
        if line.startswith('## '):
            text = line[3:].strip()
            if text == 'TABLE OF CONTENTS':
                i += 1
                continue
            add_heading1(doc, text)
            i += 1
            continue
        if line.startswith('### '):
            add_heading2(doc, line[4:].strip())
            i += 1
            continue
        if line.startswith('#### '):
            add_heading3(doc, line[5:].strip())
            i += 1
            continue

        # Blockquotes / notes
        if line.startswith('> '):
            text = line[2:].strip()
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            add_note(doc, text)
            i += 1
            continue

        # Bullets
        if re.match(r'^[-*] ', line):
            add_bullet(doc, parse_inline(line[2:]))
            i += 1
            continue
        if re.match(r'^\d+\. ', line):
            add_bullet(doc, parse_inline(re.sub(r'^\d+\.\s', '', line)))
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Horizontal rule headings (bold-only lines)
        if line.strip().startswith('---'):
            i += 1
            continue

        # Regular body
        if line.strip():
            add_body(doc, parse_inline(line))
        i += 1

    # Flush any trailing table
    if in_table and table_lines:
        h, r = parse_table(table_lines)
        if h:
            add_table(doc, h, r)

# ── Build document ────────────────────────────────────────────────────────────

doc = Document()

# Page setup — A4
sec = doc.sections[0]
sec.page_width   = Inches(8.27)
sec.page_height  = Inches(11.69)
sec.left_margin  = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.top_margin   = Inches(1.0)
sec.bottom_margin= Inches(1.0)

# Cover page
cover = doc.add_paragraph()
cover.paragraph_format.space_before = Pt(60)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run('ASPEE PHARMACEUTICALS')
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('ERP SYSTEM — STAFF USER MANUAL')
r2.bold = True; r2.font.size = Pt(18); r2.font.color.rgb = TEAL

doc.add_paragraph()

for line, sz, bold in [
    ('Version 2.0  |  April 2026',                      13, True),
    ('Prepared for: All Aspee Pharmaceuticals Staff',   11, False),
    ('Currency: Ghana Cedis (GH₵)',                11, False),
    ('CONFIDENTIAL — FOR INTERNAL USE ONLY',            10, True),
]:
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(line)
    rr.font.size = Pt(sz)
    rr.bold = bold
    rr.font.color.rgb = AMBER if bold and 'CONFIDENTIAL' in line else DGRAY

doc.add_page_break()

# Body content from markdown
MD_FILE = r'C:\Users\hp\Desktop\Developments\Aspee Pharmaceuticals\ASPEE_PHARMA_USER_MANUAL.md'
convert(MD_FILE, doc)

# Footer
doc.add_paragraph()
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = pf.add_run(
    'Aspee Pharmaceuticals ERP User Manual — Version 2.0  |  April 2026  |  '
    'Confidential — For Internal Use Only'
)
rf.font.size = Pt(9)
rf.italic = True
rf.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# Save
out = r'C:\Users\hp\Desktop\Developments\Aspee Pharmaceuticals\ASPEE_PHARMA_USER_MANUAL_v2.docx'
doc.save(out)
print('Done:', out)
