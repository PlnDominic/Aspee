from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Style helpers ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1e, 0x3a, 0x5f)
TEAL   = RGBColor(0x0f, 0x76, 0x6e)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF1, 0xF5, 0xF9)
DGRAY  = RGBColor(0x33, 0x33, 0x33)
AMBER  = RGBColor(0x92, 0x40, 0x0E)
GREEN  = RGBColor(0x16, 0x4e, 0x32)

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        val = kwargs.get(edge, {})
        if val:
            node = OxmlElement(f'w:{edge}')
            for k,v in val.items():
                node.set(qn(f'w:{k}'), v)
            tcBorders.append(node)
    tcPr.append(tcBorders)

def para_space(doc, before=0, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    return p

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # left bar via shading trick — use a styled run
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(15)
    run.font.color.rgb = NAVY
    # bottom border on paragraph
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '12')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1e3a5f')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = TEAL
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = DGRAY
    return p

def body(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DGRAY
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def note_box(doc, text):
    """Shaded note paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'),   'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'),  'FEF3C7')
    p._p.get_or_add_pPr().append(shading_elm)
    run = p.add_run('  Note: ' + text)
    run.font.size  = Pt(9.5)
    run.font.color.rgb = AMBER
    run.italic = True
    return p

def make_table(doc, headers, rows, col_widths=None):
    """Build a styled table with a navy header row."""
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '1e3a5f')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = 'F1F5F9' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p   = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def page_break(doc):
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.paragraph_format.space_before = Pt(60)
cover.paragraph_format.space_after  = Pt(4)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run('ASPEE PHARMACEUTICALS')
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('ERP SYSTEM — STAFF USER MANUAL')
r2.bold = True; r2.font.size = Pt(18); r2.font.color.rgb = TEAL

doc.add_paragraph().paragraph_format.space_after = Pt(30)

for line, sz in [
    ('Version 1.0  |  April 2026', 12),
    ('Prepared for: All Aspee Pharmaceuticals Staff', 11),
    ('Currency: Ghana Cedis (GH₵)', 11),
    ('CONFIDENTIAL — FOR INTERNAL USE ONLY', 10),
]:
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(line)
    rr.font.size = Pt(sz)
    rr.font.color.rgb = DGRAY
    if 'CONFIDENTIAL' in line:
        rr.bold = True; rr.font.color.rgb = AMBER

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (static — Word will update field on open)
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, 'Table of Contents')
toc_items = [
    ('1.', 'Getting Started — Login & Navigation'),
    ('2.', 'Role & Access Guide'),
    ('3.', 'Dashboard Overview'),
    ('4.', 'Procurement Department'),
    ('5.', 'Quality Assurance (QA)'),
    ('6.', 'Stores Department'),
    ('7.', 'Production Department'),
    ('8.', 'Sales Department'),
    ('9.', 'Accounting Department'),
    ('10.', 'Human Resources (HR)'),
    ('11.', 'Internal Audit'),
    ('12.', 'Compliance & Regulators'),
    ('13.', 'Weekly Reports'),
    ('14.', 'Settings & Administration (Super Admin)'),
    ('15.', 'Company-Wide Workflow Summary'),
    ('16.', 'Common Tasks Quick Reference'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{num}  ')
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
    r2 = p.add_run(title)
    r2.font.size = Pt(10); r2.font.color.rgb = DGRAY

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — GETTING STARTED
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '1.  Getting Started — Login & Navigation')

heading2(doc, '1.1  Accessing the System')
body(doc, 'Open your web browser (Chrome or Edge recommended) and navigate to the Aspee Pharma ERP URL provided by your IT administrator. You will land on the Login Page.')

heading2(doc, '1.2  Logging In')
bullet(doc, 'Enter your company email address and password, then click Sign In.')
bullet(doc, 'If you forget your password, contact your Super Admin to reset it.')
note_box(doc, 'If your account does not yet exist in the system, ask your Super Admin to add you under Settings → User Management.')

heading2(doc, '1.3  The Sidebar Navigation')
body(doc, 'After login you will see the left sidebar — your main navigation panel.')
bullet(doc, 'Click any section heading (e.g., Procurement, Sales) to expand its sub-menu.')
bullet(doc, 'The active page is highlighted in blue.')
bullet(doc, 'Click the Collapse button at the bottom to hide labels and save screen space.')
bullet(doc, 'The sidebar only shows sections your role is permitted to access.')

heading2(doc, '1.4  Common Page Actions')
make_table(doc,
    ['Element', 'Purpose'],
    [
        ['Page Header', 'Title, subtitle, and breadcrumb trail showing where you are'],
        ['Stat Cards', '4 key metrics shown at the top of each section'],
        ['Data Table', 'Searchable list of all records'],
        ['+ New / Create button', 'Opens the form to add a new record'],
        ['Eye (View) icon', 'Opens a read-only detail view'],
        ['Pencil (Edit) icon', 'Opens the record for editing'],
        ['Trash (Delete) icon', 'Deletes the record (with confirmation prompt)'],
        ['Export button', 'Downloads current data as a CSV file'],
        ['Send Weekly Report button', 'Sends a departmental summary to the Managing Director'],
    ],
    col_widths=[2.2, 4.3]
)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — ROLES
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '2.  Role & Access Guide')
body(doc, 'Each staff member is assigned exactly one role. The system automatically shows only the modules permitted for that role.')

make_table(doc,
    ['Role', 'Accessible Modules'],
    [
        ['Super Admin', 'Everything — full system access'],
        ['Managing Director', 'Dashboard, Weekly Reports Review'],
        ['Sales Manager', 'Sales (all), Customers, Weekly Reports'],
        ['Van Sales Rep', 'Invoices, Dispatch, Receipts, Weekly Reports'],
        ['Purchasing Manager', 'Procurement, Supplier Payments, GRN, Weekly Reports'],
        ['Store Manager', 'Stores (all), Production, Weekly Reports'],
        ['Production Manager', 'Production, Weekly Reports'],
        ['Quality Assurance', 'QA (all), GRN, Compliance, Weekly Reports'],
        ['Accountant', 'Accounting (all), Collections, Supplier Payments, Weekly Reports'],
        ['HR Manager', 'HR (all), Weekly Reports'],
        ['Internal Auditor', 'Internal Audit, Weekly Reports'],
    ],
    col_widths=[2.0, 4.5]
)
note_box(doc, 'If you try to access a page outside your role, the system will automatically redirect you to the Dashboard.')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — DASHBOARD
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '3.  Dashboard Overview')
body(doc, 'Path: Overview (Home Page after Login)')
body(doc, 'The dashboard gives every staff member a bird\'s-eye view of company performance. It displays summary cards, quick links to common tasks, and alerts for low stock, pending approvals, or overdue invoices.')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — PROCUREMENT
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '4.  Procurement Department')
body(doc, 'Accessible to: Super Admin, Purchasing Manager      |      Sidebar Path: Procurement')

heading2(doc, '4.1  Suppliers')
body(doc, 'Path: Procurement → Suppliers')
body(doc, 'The Suppliers list is the master record of all vendors Aspee Pharmaceuticals does business with.')
body(doc, 'To add a new supplier: click + Add Supplier, fill in the fields below, then click Save.')

make_table(doc,
    ['Field', 'Example Value'],
    [
        ['Supplier Name',   'Meridian Pharma Supplies, Tema'],
        ['Contact Person',  'Mr. Kwame Asante'],
        ['Phone',           '0244 123 456'],
        ['Email',           'orders@meridianpharma.com'],
        ['Address',         'Tema Industrial Area, Block 5'],
        ['Payment Terms',   'Net 30 days'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '4.2  Purchase Orders (POs)')
body(doc, 'Path: Procurement → Purchase Orders')
body(doc, 'A Purchase Order is a formal request sent to a supplier to purchase goods.')

heading3(doc, 'Stat Cards on this page')
make_table(doc,
    ['Stat Card', 'Meaning'],
    [
        ['Total POs',      'Number of all purchase orders in the system'],
        ['Pending Review', 'POs awaiting approval'],
        ['Completed',      'POs that have been fully received'],
        ['Total Value',    'Combined GH₵ value of all POs'],
    ],
    col_widths=[2.2, 4.3]
)

heading3(doc, 'PO Status Lifecycle')
body(doc, 'Pending  →  Approved  →  Shipped  →  Received  →  Closed')
body(doc, '                                                 ↓')
body(doc, '                                  Cancelled (if rejected)')

heading3(doc, 'Approval Rules')
bullet(doc, 'POs under GH₵ 10,000  →  Approved at Manager level.')
bullet(doc, 'POs over GH₵ 10,000   →  Require Finance (Accountant) approval.')
bullet(doc, 'Click the green checkmark (✓) icon on any Pending PO to approve it.')

heading3(doc, 'Example Purchase Order')
make_table(doc,
    ['Field', 'Value'],
    [
        ['PO Number',       'PO-2026-047'],
        ['Supplier',        'Meridian Pharma Supplies, Tema'],
        ['Date',            '10 April 2026'],
        ['Item 1',          'Paracetamol API — 500 kg @ GH₵ 120/kg = GH₵ 60,000'],
        ['Item 2',          'Maize Starch — 200 kg @ GH₵ 45/kg = GH₵ 9,000'],
        ['Total Amount',    'GH₵ 69,000.00'],
        ['Approval Level',  'Finance (amount exceeds GH₵ 10,000)'],
        ['Status',          'Approved'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '4.3  Requisitions (Purchase Requests)')
body(doc, 'Path: Procurement → Requisitions')
body(doc, 'Internal requests from Stores or Production asking Procurement to purchase goods. The Purchasing Manager reviews these and converts approved requests into formal Purchase Orders.')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — QA
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '5.  Quality Assurance (QA)')
body(doc, 'Accessible to: Super Admin, Quality Assurance Officer      |      Sidebar Path: Quality Assurance')
body(doc, 'QA is the gatekeeper of all materials and finished products. Nothing enters Stores without QA sign-off.')

heading2(doc, '5.1  QA Overview')
body(doc, 'Path: Quality Assurance → Overview')
body(doc, 'A dashboard showing the status of all ongoing QA activities — incoming inspections, in-process checks, and finished product releases.')

heading2(doc, '5.2  Incoming Materials Inspection')
body(doc, 'Path: Quality Assurance → Incoming Materials')
body(doc, 'When a supplier delivers goods, QA must inspect them before they are accepted into the warehouse. The result is one of three statuses: Approved, Rejected, or Quarantine.')

heading3(doc, 'Example Incoming Inspection Record')
make_table(doc,
    ['Field', 'Value'],
    [
        ['GRN Number',       'GRN-2026-022'],
        ['Supplier',         'Meridian Pharma Supplies'],
        ['Product',          'Paracetamol API'],
        ['Batch Number',     'BATCH-PAR-0041'],
        ['Qty Ordered',      '500 kg'],
        ['Qty Received',     '498 kg'],
        ['Appearance',       'White crystalline powder — Pass'],
        ['Assay Result',     '99.2% — Pass'],
        ['QA Result',        'Approved'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '5.3  Goods Receipt Note (GRN)')
body(doc, 'Path: Quality Assurance → Goods Receipt')
body(doc, 'The GRN is the official record that goods have been received. It must be linked to the original Purchase Order. After QA approves GRN items, the approved quantity is added to Stock Inventory.')

heading2(doc, '5.4  In-Process Controls')
body(doc, 'Path: Quality Assurance → In Process Controls')
body(doc, 'During manufacturing, QA performs checks at defined production stages to ensure the product is being made correctly.')

heading3(doc, 'Example In-Process Check — Tablet Compression')
make_table(doc,
    ['Check Point', 'Parameter', 'Acceptable Range', 'Result'],
    [
        ['Weight',        'Average tablet weight',  '500 mg ± 5%',  '502 mg — Pass'],
        ['Hardness',      'Tablet hardness',        '4–8 kp',       '6.2 kp — Pass'],
        ['Disintegration','Disintegration time',    '< 15 minutes', '8 min — Pass'],
        ['Appearance',    'Visual check',           'No defects',   'Pass'],
    ],
    col_widths=[1.6, 1.8, 1.7, 1.4]
)

heading2(doc, '5.5  Finished Products Analysis')
body(doc, 'Path: Quality Assurance → Finished Products')
body(doc, 'Before any manufactured batch can be released for sale, it must pass QA\'s final analysis.')
bullet(doc, 'Passed — Product released to Stores (finished goods stock).')
bullet(doc, 'Failed — Product rejected and destroyed or quarantined.')
bullet(doc, 'Quarantine — Under investigation, awaiting final decision.')
body(doc, 'Quick Approve: For batches in Quarantine that have subsequently cleared review, click the green "Approve & Release" button directly from the table.')

heading3(doc, 'Example Finished Products Record')
make_table(doc,
    ['Field', 'Value'],
    [
        ['Product Name',    'Paracetamol 500mg Tablets'],
        ['Batch Number',    'BPC-2026-018'],
        ['Job Order',       'JO-2026-011'],
        ['Analyst',         'Abena Mensah'],
        ['Analysis Date',   '05 April 2026'],
        ['Dissolution',     '85% at 30 min — Pass'],
        ['Assay',           '98.8% — Pass'],
        ['Microbial (TAMC)','< 1,000 cfu/g — Pass'],
        ['Overall Status',  'Passed'],
        ['Release Date',    '07 April 2026'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '5.6  Internal QA Reports')
body(doc, 'Path: Quality Assurance → Internal Reports')
body(doc, 'QA internal deviation reports, non-conformance records, and corrective action summaries are stored here.')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — STORES
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '6.  Stores Department')
body(doc, 'Accessible to: Super Admin, Store Manager      |      Sidebar Path: Stores')
body(doc, 'The Stores department is the physical custodian of all raw materials, packaging materials, and finished goods.')

heading2(doc, '6.1  Products (Product Master)')
body(doc, 'Path: Stores → Products')
body(doc, 'The master list of every product or material the company handles. All items must be registered here before they can be stocked or sold.')

make_table(doc,
    ['Field', 'Description', 'Example'],
    [
        ['Product Name',   'Full name of the item',                    'Paracetamol 500mg Tablets'],
        ['SKU',            'Unique stock-keeping code',                 'PAR-500-TAB'],
        ['Unit',           'Unit of measurement',                       'Bottles / Kg / Boxes'],
        ['Material Type',  'Raw Material / Finished Product / Packaging','Finished Product'],
        ['Reorder Level',  'Minimum stock before alert triggers',        '500'],
    ],
    col_widths=[1.8, 2.4, 2.3]
)
note_box(doc, 'The reorder level is critical. When stock drops to or below this number, the system automatically alerts staff with a low-stock warning.')

heading2(doc, '6.2  Stock Inventory')
body(doc, 'Path: Stores → Stock Inventory')
body(doc, 'Real-time view of stock across all locations — main warehouse, production floor, sales vans, etc.')

make_table(doc,
    ['Status', 'Meaning'],
    [
        ['Adequate',     'Stock is comfortably above the reorder level'],
        ['Warning',      'Stock is between 1× and 1.5× the reorder level'],
        ['Low',          'Stock is at or below the reorder level'],
        ['Out of Stock', 'Zero quantity on hand'],
    ],
    col_widths=[2.0, 4.5]
)

heading3(doc, 'Example Stock Inventory Table')
make_table(doc,
    ['Product', 'SKU', 'Unit', 'Warehouse', 'Van 1', 'Total', 'Reorder', 'Status'],
    [
        ['Paracetamol 500mg', 'PAR-500-TAB', 'Bottles', '4,200', '300', '4,500', '1,000', 'Adequate'],
        ['Amoxicillin 250mg', 'AMX-250-CAP', 'Bottles',   '180', '120',   '300',   '500', 'Low'],
        ['Vitamin C 200mg',   'VIT-C-200',   'Bottles',     '0',   '0',     '0',   '200', 'Out of Stock'],
    ],
    col_widths=[1.6, 1.1, 0.7, 0.9, 0.6, 0.6, 0.7, 0.8]
)

heading2(doc, '6.3  Stock Transfers')
body(doc, 'Path: Stores → Stock Transfers')
body(doc, 'Used to move stock between locations — e.g., from the Warehouse to a Sales Van, or from the Warehouse to the Production Floor.')

heading3(doc, 'Example Transfer')
make_table(doc,
    ['Field', 'Value'],
    [
        ['Transfer No',  'TRF-2026-034'],
        ['From',         'Main Warehouse'],
        ['To',           'Van 1 (Kwaku Asante — Accra Central Route)'],
        ['Product',      'Paracetamol 500mg Tablets — 300 bottles'],
        ['Date',         '10 April 2026'],
    ],
    col_widths=[2.2, 4.3]
)
note_box(doc, 'Stores must load stock to a sales van before the Van Sales Rep can issue invoices. The system will block invoice creation if the van has no loaded stock.')

heading2(doc, '6.4  Material Requests')
body(doc, 'Path: Stores → Material Requests')
body(doc, 'Production raises Material Requests when they need raw materials from Stores for a manufacturing job. The Store Manager reviews, issues the materials, and records the stock movement (OUT).')

heading2(doc, '6.5  Purchase Requests')
body(doc, 'Path: Stores → Purchase Requests')
body(doc, 'When a material is needed but not in stock, the Store Manager raises a Purchase Request. This goes to the Purchasing Manager to create a formal Purchase Order.')

heading2(doc, '6.6  Sales Movements')
body(doc, 'Path: Stores → Sales Movements')
body(doc, 'A read-only view of all stock movements caused by sales invoices — showing which products left which van, when, and against which invoice.')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — PRODUCTION
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '7.  Production Department')
body(doc, 'Accessible to: Super Admin, Production Manager, Store Manager      |      Sidebar Path: Production')

heading2(doc, '7.1  Bill of Materials (BOM)')
body(doc, 'Path: Production → Bill of Materials')
body(doc, 'The BOM defines the exact recipe for producing a finished product — which raw materials are needed, in what quantities, and in what units.')

heading3(doc, 'Example BOM — Paracetamol 500mg Tablets (Batch of 10,000 tablets)')
make_table(doc,
    ['Component', 'Quantity', 'Unit'],
    [
        ['Paracetamol API',              '5.00', 'kg'],
        ['Maize Starch',                 '1.20', 'kg'],
        ['Microcrystalline Cellulose',   '0.80', 'kg'],
        ['Magnesium Stearate',           '0.05', 'kg'],
        ['Talc',                         '0.10', 'kg'],
        ['PVC Blister Foil',            '200',   'sheets'],
    ],
    col_widths=[3.5, 1.5, 1.5]
)

heading2(doc, '7.2  Job Orders (Production Orders)')
body(doc, 'Path: Production → Job Orders')
body(doc, 'A Job Order is the formal instruction to manufacture a specific batch of a product.')
body(doc, 'Job Order Status:  Planned  →  In Progress  →  Completed  →  Released (after QA Approval)')

heading3(doc, 'Example Job Order')
make_table(doc,
    ['Field', 'Value'],
    [
        ['Order Number',   'JO-2026-015'],
        ['Product',        'Paracetamol 500mg Tablets'],
        ['Batch Size',     '50,000 tablets'],
        ['Start Date',     '08 April 2026'],
        ['Expected End',   '12 April 2026'],
        ['Status',         'In Progress'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '7.3  Production Material Requests')
body(doc, 'Path: Production → Material Requests')
body(doc, 'When Production starts a Job Order, they must request the required raw materials from Stores.')
bullet(doc, 'Production Manager creates a Material Request linked to the Job Order.')
bullet(doc, 'Store Manager reviews the request and checks stock availability.')
bullet(doc, 'If available, Stores issues the materials — stock quantity decreases.')
bullet(doc, 'Completed batch goes to QA for Finished Products Analysis.')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — SALES
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '8.  Sales Department')
body(doc, 'Accessible to: Super Admin, Sales Manager, Van Sales Rep      |      Sidebar Path: Sales')

heading2(doc, '8.1  Customers')
body(doc, 'Path: Sales → Customers')
body(doc, 'All customers must be registered here before an invoice can be raised for them.')

make_table(doc,
    ['Category', 'Description'],
    [
        ['Wholesale',   'Large bulk buyers'],
        ['Retail',      'Direct retail outlets'],
        ['Hospital',    'Government or private hospitals'],
        ['Pharmacy',    'Registered pharmacy shops'],
        ['Institution', 'Schools, NGOs, government agencies'],
    ],
    col_widths=[2.0, 4.5]
)

heading3(doc, 'Example Customer Record')
make_table(doc,
    ['Field', 'Value'],
    [
        ['Customer Name',   'Kumasi Central Pharmacy'],
        ['Contact Person',  'Mrs. Abena Osei'],
        ['Phone',           '0322 456 789'],
        ['Category',        'Pharmacy'],
        ['Credit Limit',    'GH₵ 15,000'],
        ['Payment Terms',   'Net 14 days'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '8.2  Routes & Vans')
body(doc, 'Path: Sales → Routes & Vans')
body(doc, 'Each Sales Van is assigned to a specific delivery route/territory. The van\'s route determines which stock location is deducted when invoices are issued.')

make_table(doc,
    ['Field', 'Value'],
    [
        ['Van ID',       'VAN-001'],
        ['Driver',       'Kwaku Asante'],
        ['Plate Number', 'GR 1234-23'],
        ['Route Area',   'Accra Central'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '8.3  Sales Invoices')
body(doc, 'Path: Sales → Invoices')

make_table(doc,
    ['Invoice Type', 'Description'],
    [
        ['Cash Sale',   'Payment received immediately at point of sale.'],
        ['Credit Sale', 'Payment due on a future date (based on customer credit terms).'],
    ],
    col_widths=[2.0, 4.5]
)

body(doc, 'Invoice Status Lifecycle:')
body(doc, 'Draft  →  Issued  →  Partially Paid  →  Paid')
body(doc, '                          ↓')
body(doc, '                    Overdue (if payment due date passes without full collection)')

note_box(doc, 'Only Draft invoices can be deleted. Once an invoice is Issued, it can only be reversed via a Credit Note.')
note_box(doc, 'When an invoice is saved as Issued, the system automatically: (1) deducts sold quantities from the van\'s stock, (2) records a stock movement (OUT), (3) posts an automatic journal entry, and (4) alerts if any product drops below its reorder level.')

heading3(doc, 'Example Sales Invoice')
make_table(doc,
    ['Field', 'Value'],
    [
        ['Invoice No',    'INV-2026-0247'],
        ['Customer',      'Kumasi Central Pharmacy'],
        ['Date',          '10 April 2026'],
        ['Due Date',      '24 April 2026 (Net 14 days)'],
        ['Type',          'Credit Sale'],
        ['Van / Route',   'VAN-001 (Accra Central — Kwaku Asante)'],
        ['Item 1',        'Paracetamol 500mg Tablets — 50 bottles @ GH₵ 48.00 = GH₵ 2,400.00'],
        ['Item 2',        'Vitamin C 200mg Tablets — 20 bottles @ GH₵ 35.00 = GH₵ 700.00'],
        ['Total Amount',  'GH₵ 3,100.00'],
        ['Status',        'Issued'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '8.4  Dispatch Management')
body(doc, 'Path: Sales → Dispatch Management')
body(doc, 'Dispatch records track the physical delivery of goods to customers.')
body(doc, 'Status: Draft  →  Pending  →  In Transit  →  Completed (or Cancelled)')

heading3(doc, 'Example Dispatch Record')
make_table(doc,
    ['Field', 'Value'],
    [
        ['Dispatch No',         'DSP-2026-089'],
        ['Van',                 'VAN-001 (GR 1234-23)'],
        ['Driver',              'Kwaku Asante'],
        ['Dispatch Date',       '10 April 2026'],
        ['Invoices Included',   'INV-2026-0247, INV-2026-0248, INV-2026-0249'],
        ['Status',              'In Transit'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '8.5  Receipts (Payment Collection)')
body(doc, 'Path: Sales → Receipts')
body(doc, 'Receipts record payments collected from customers against their invoices.')
note_box(doc, 'Confirmed or Cleared receipts cannot be deleted — contact the Super Admin to void if recorded in error.')

heading3(doc, 'Example Receipt')
make_table(doc,
    ['Field', 'Value'],
    [
        ['Receipt No',       'RCP-2026-0183'],
        ['Invoice',          'INV-2026-0247'],
        ['Customer',         'Kumasi Central Pharmacy'],
        ['Amount Collected', 'GH₵ 3,100.00'],
        ['Payment Method',   'Mobile Money (MoMo)'],
        ['MoMo Reference',   'GH-MM-739281'],
        ['Date Collected',   '17 April 2026'],
        ['Status',           'Confirmed'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '8.6  Credit Notes')
body(doc, 'Path: Sales → Credit Notes')
body(doc, 'A Credit Note is issued when a customer returns goods or when an invoice needs to be partially or fully reversed.')
bullet(doc, 'Customer returns defective products.')
bullet(doc, 'Over-billing error on an issued invoice.')
bullet(doc, 'Agreed price adjustment after invoice issuance.')

heading3(doc, 'Example Credit Note')
make_table(doc,
    ['Field', 'Value'],
    [
        ['Credit Note No',     'CN-2026-0021'],
        ['Original Invoice',   'INV-2026-0247'],
        ['Customer',           'Kumasi Central Pharmacy'],
        ['Reason',             '5 bottles of Paracetamol returned (expired stock)'],
        ['Credit Amount',      'GH₵ 240.00'],
        ['Date',               '18 April 2026'],
    ],
    col_widths=[2.2, 4.3]
)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — ACCOUNTING
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '9.  Accounting Department')
body(doc, 'Accessible to: Super Admin, Accountant      |      Sidebar Path: Accounting')

heading2(doc, '9.1  Journal Entries')
body(doc, 'Path: Accounting → Journal Entries')
body(doc, 'Manual journal entries for transactions not automatically posted by the system. Most sales transactions auto-post when invoices are issued.')

make_table(doc,
    ['Account', 'Type', 'Normal Balance'],
    [
        ['Sales Revenue',     'Income',    'Credit'],
        ['Accounts Receivable','Asset',    'Debit'],
        ['Cash & Bank',       'Asset',     'Debit'],
        ['Inventory / Stock', 'Asset',     'Debit'],
        ['Cost of Goods Sold','Expense',   'Debit'],
        ['Accounts Payable',  'Liability', 'Credit'],
        ['VAT Payable',       'Liability', 'Credit'],
    ],
    col_widths=[2.5, 1.5, 2.5]
)

heading2(doc, '9.2  General Ledger')
body(doc, 'Path: Accounting → General Ledger')
body(doc, 'All posted journal entries organised by account. Use this to trace any transaction back to its origin.')

heading2(doc, '9.3  Trial Balance')
body(doc, 'Path: Accounting → Trial Balance')
body(doc, 'A summary of all account balances confirming total debits equal total credits. Run at month-end.')

heading2(doc, '9.4  Expenses')
body(doc, 'Path: Accounting → Expenses')

heading3(doc, 'Example Expense Entry')
make_table(doc,
    ['Field', 'Value'],
    [
        ['Category',        'Fuel & Transport'],
        ['Description',     'Fuel for Van 1 delivery route — April 2026'],
        ['Amount',          'GH₵ 850.00'],
        ['Date',            '10 April 2026'],
        ['Payment Method',  'Cash (Petty Cash)'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '9.5  Payroll')
body(doc, 'Path: Accounting → Payroll')

heading3(doc, 'Example Payroll Entries')
make_table(doc,
    ['Employee', 'Basic (GH₵)', 'Transport (GH₵)', 'SSNIT 5.5% (GH₵)', 'Tax (GH₵)', 'Net Pay (GH₵)'],
    [
        ['Abena Mensah (QA Analyst)',   '3,800', '400', '209.00', '312.00', '3,679.00'],
        ['Kwaku Asante (Van Driver)',   '2,500', '300', '137.50', '178.00', '2,484.50'],
        ['Kofi Darko (Store Keeper)',   '2,200', '250', '121.00', '148.00', '2,181.00'],
    ],
    col_widths=[1.9, 1.0, 1.1, 1.2, 0.9, 1.2]
)

heading2(doc, '9.6  Tax Periods')
body(doc, 'Path: Accounting → Tax Periods')
body(doc, 'Track VAT filing periods, corporate tax assessments, and PAYE submissions to GRA.')

heading2(doc, '9.7  Petty Cash')
body(doc, 'Path: Accounting → Petty Cash')

make_table(doc,
    ['Date', 'Description', 'Amount (GH₵)', 'Balance (GH₵)'],
    [
        ['07 Apr 2026', 'Office stationery — pens, paper', '85.00',  '415.00'],
        ['08 Apr 2026', 'Tea/coffee for factory floor',    '45.00',  '370.00'],
        ['10 Apr 2026', 'Courier delivery fee',            '30.00',  '340.00'],
    ],
    col_widths=[1.4, 2.8, 1.5, 1.5]
)

heading2(doc, '9.8  Supplier Payments')
body(doc, 'Path: Accounting → Supplier Payments')

make_table(doc,
    ['Field', 'Value'],
    [
        ['Supplier',        'Meridian Pharma Supplies'],
        ['PO Reference',    'PO-2026-047'],
        ['Invoice Amount',  'GH₵ 69,000.00'],
        ['Amount Paid',     'GH₵ 34,500.00 (50% deposit)'],
        ['Balance',         'GH₵ 34,500.00'],
        ['Payment Method',  'Bank Transfer'],
        ['Bank Reference',  'GCB-TRF-20260408'],
        ['Date Paid',       '08 April 2026'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '9.9  A/R Aging Report')
body(doc, 'Path: Accounting → A/R Aging')
body(doc, 'Shows all outstanding customer invoices grouped by how long they have been unpaid.')

make_table(doc,
    ['Aging Bucket', 'Meaning', 'Action Required'],
    [
        ['0–30 Days',   'Current — recently issued',       'Monitor weekly'],
        ['31–60 Days',  'Slightly overdue',                'Send payment reminder'],
        ['61–90 Days',  'Significantly overdue',           'Escalate to Sales Manager'],
        ['90+ Days',    'Seriously overdue',               'Legal / MD escalation'],
    ],
    col_widths=[1.4, 2.3, 2.8]
)

heading3(doc, 'Example A/R Aging')
make_table(doc,
    ['Customer', 'Invoice', 'Total (GH₵)', 'Paid (GH₵)', 'Outstanding (GH₵)', 'Bucket'],
    [
        ['Kumasi Central Pharmacy', 'INV-2026-0187', '5,200.00', '2,000.00', '3,200.00', '31–60 Days'],
        ['Accra Medical Stores',    'INV-2026-0162', '8,400.00', '0.00',     '8,400.00', '61–90 Days'],
    ],
    col_widths=[1.7, 1.3, 1.2, 1.1, 1.4, 1.0]
)

heading2(doc, '9.10  Financial Statements')
make_table(doc,
    ['Statement', 'Path', 'Purpose'],
    [
        ['Comprehensive Income',    'Accounting → Comprehensive Income',    'Revenues, costs, profit/loss'],
        ['Financial Position',      'Accounting → Financial Position',      'Balance sheet — assets, liabilities, equity'],
        ['Equity Statement',        'Accounting → Equity',                  'Changes in owners\' equity'],
        ['Cash Flow Statement',     'Accounting → Cash Flow',               'Cash inflows and outflows'],
        ['Accounting Notes',        'Accounting → Accounting Notes',        'Disclosure notes'],
    ],
    col_widths=[1.8, 2.2, 2.5]
)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — HR
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '10.  Human Resources (HR)')
body(doc, 'Accessible to: Super Admin, HR Manager      |      Sidebar Path: HR Management')

heading2(doc, '10.1  Employees')
body(doc, 'Path: HR → Employees')
make_table(doc,
    ['Field', 'Example Value'],
    [
        ['Employee ID',      'EMP-2024-018'],
        ['Full Name',        'Abena Mensah'],
        ['Department',       'Quality Assurance'],
        ['Job Title',        'QA Analyst'],
        ['Date Employed',    '15 June 2024'],
        ['Phone',            '0244 987 654'],
        ['Emergency Contact','Mr. Kofi Mensah (Husband) — 0201 123 456'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '10.2  Attendance')
body(doc, 'Path: HR → Attendance')
body(doc, 'Daily attendance records including check-in/check-out times, hours worked, and absent/present/late status.')

heading2(doc, '10.3  Leave Management')
body(doc, 'Path: HR → Leave Management')
body(doc, 'Track and approve staff leave requests. Leave types include Annual, Sick, Maternity/Paternity, and Compassionate leave.')
bullet(doc, 'Employee (or HR on their behalf) submits a leave request.')
bullet(doc, 'HR Manager reviews and approves or rejects.')
bullet(doc, 'Approved leave updates the employee\'s leave balance.')

heading2(doc, '10.4  Payroll Preparation')
body(doc, 'Path: HR → Payroll Preparation')
body(doc, 'HR prepares payroll data (attendance records, leave deductions, overtime) which is then processed by Accounting.')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — INTERNAL AUDIT
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '11.  Internal Audit')
body(doc, 'Accessible to: Super Admin, Internal Auditor      |      Sidebar Path: Internal Audit')

heading2(doc, '11.1  Audit Plans')
body(doc, 'Path: Internal Audit → Audit Plans')

make_table(doc,
    ['Field', 'Value'],
    [
        ['Audit Title',   'Q2 2026 Stores Inventory Audit'],
        ['Department',    'Stores'],
        ['Planned Date',  '30 April 2026'],
        ['Auditor',       'Emmanuel Darko (Internal Auditor)'],
        ['Scope',         'Reconcile physical stock count against ERP inventory records'],
    ],
    col_widths=[2.2, 4.3]
)

heading2(doc, '11.2  Audit Reports')
body(doc, 'Path: Internal Audit → Audit Reports')
body(doc, 'Records findings, recommendations, management responses, and target resolution dates from completed audits.')

heading2(doc, '11.3  Non-Conformances')
body(doc, 'Path: Internal Audit → Non-Conformances')
body(doc, 'Any deviation from established procedures, standards, or regulations.')

make_table(doc,
    ['Severity', 'Meaning'],
    [
        ['Minor',    'Isolated, low-risk deviation'],
        ['Major',    'Systematic or high-risk deviation'],
        ['Critical', 'Immediate risk to product quality, patient safety, or regulatory compliance'],
    ],
    col_widths=[1.5, 5.0]
)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 12 — COMPLIANCE
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '12.  Compliance & Regulators')
body(doc, 'Accessible to: Super Admin, Quality Assurance, Managing Director      |      Path: Compliance → Regulators & Renewals')
body(doc, 'Track all regulatory licences, certifications, and renewal dates to ensure Aspee Pharmaceuticals remains compliant with Ghana FDA and other regulatory bodies.')

make_table(doc,
    ['Regulator', 'Licence Type', 'Licence No', 'Expiry', 'Status'],
    [
        ['FDA Ghana', 'Manufacturer\'s Licence', 'MFR-GH-2024-0089', '31 Dec 2026', 'Active'],
        ['EPA',       'Factory Registration',    'EPA-2023-FAC-1102','30 Jun 2026', 'Due for Renewal'],
        ['GSS',       'Business Registration',   'GSS-REG-20190',    'Permanent',   'Active'],
    ],
    col_widths=[1.2, 1.8, 1.8, 1.3, 1.4]
)
note_box(doc, 'Set calendar reminders 90 days before any expiry date to initiate renewal with the relevant authority.')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 13 — WEEKLY REPORTS
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '13.  Weekly Reports')
body(doc, 'Accessible to: All staff (submission)      |      Review: Super Admin and Managing Director only')

heading2(doc, '13.1  Submitting a Weekly Report')
body(doc, 'Each department head submits a weekly summary to the Managing Director.')
bullet(doc, 'Go to the relevant module page (e.g., Stores → Stock Inventory).')
bullet(doc, 'Click the teal "Send Weekly Report" button in the top-right of the page.')
bullet(doc, 'Confirm the submission — it will be delivered to the MD\'s review inbox.')

heading2(doc, '13.2  Reviewing Weekly Reports (Managing Director)')
body(doc, 'Path: Weekly Reports → Review')
body(doc, 'The Managing Director reviews all submitted department reports in one consolidated view.')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 14 — SETTINGS
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '14.  Settings & Administration')
body(doc, 'Accessible to: Super Admin only (except Profile which is accessible to all)')

heading2(doc, '14.1  User Management')
body(doc, 'Path: Settings → User Management')
body(doc, 'Create system users and assign roles. To add a new user: click + Add User, enter the staff email, select their role, and save.')

make_table(doc,
    ['Role', 'Description'],
    [
        ['Super Admin',         'Full access — IT Administrator only'],
        ['Managing Director',   'Executive overview and report review'],
        ['Sales Manager',       'Manages sales team and all sales data'],
        ['Van Sales Rep',       'Issues invoices and records deliveries on the road'],
        ['Purchasing Manager',  'Manages procurement and supplier relationships'],
        ['Store Manager',       'Controls inventory and stock movements'],
        ['Production Manager',  'Oversees manufacturing job orders'],
        ['Quality Assurance',   'Inspects materials and approves product releases'],
        ['Accountant',          'Manages all financial records and statements'],
        ['HR Manager',          'Manages staff records, attendance, and leave'],
        ['Internal Auditor',    'Conducts internal audits across all departments'],
    ],
    col_widths=[2.0, 4.5]
)

heading2(doc, '14.2  Audit Trail')
body(doc, 'Path: Settings → Audit Trail')
body(doc, 'A complete immutable log of every action in the system — who did it, what they did, which record, and when. All create, edit, and delete actions are recorded.')

heading2(doc, '14.3  Report Settings')
body(doc, 'Path: Settings → Report Settings')
body(doc, 'Configure company name, address, logo, and other details that appear on printed invoices and reports.')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 15 — WORKFLOW SUMMARY
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '15.  Company-Wide Workflow Summary')
body(doc, 'The Aspee Pharmaceuticals ERP follows a linear, interconnected workflow across all departments. Each step feeds into the next.')

steps = [
    ('1. PROCUREMENT',
     'Store Manager raises a Purchase Request  →  Purchasing Manager creates a Purchase Order  →  Supplier delivers goods.'),
    ('2. QUALITY ASSURANCE — INCOMING',
     'QA Officer inspects the delivery and records a GRN  →  QA approves, rejects, or quarantines materials  →  Approved materials are added to Warehouse stock.'),
    ('3. STORES',
     'Store Manager oversees physical stock  →  Production raises a Material Request  →  Stores issues materials (stock decreases).'),
    ('4. PRODUCTION',
     'Production Manager creates a Job Order linked to the BOM  →  Factory manufactures the batch  →  Finished goods sent to QA for final testing.'),
    ('5. QUALITY ASSURANCE — FINISHED PRODUCTS',
     'QA Analyst records Finished Products Analysis  →  Batch Passed / Failed / Quarantined  →  Passed batches released to Finished Goods stock.'),
    ('6. STORES — OUTBOUND',
     'Stores loads stock onto the Sales Van via a Stock Transfer  →  Van is now ready for sales.'),
    ('7. SALES',
     'Van Sales Rep issues a Sales Invoice (stock deducted from van)  →  Dispatch record created and updated to Completed  →  Customer pays — Receipt recorded  →  Returns handled via Credit Notes.'),
    ('8. ACCOUNTING',
     'All invoices auto-post journal entries  →  Accountant manages expenses, payroll, supplier payments  →  Weekly: A/R Aging reviewed  →  Monthly: Trial Balance and Financial Statements generated.'),
    ('9. INTERNAL AUDIT & COMPLIANCE',
     'Regular audits across all departments  →  Non-conformances investigated and closed  →  Regulatory licences monitored and renewed.'),
    ('10. MANAGING DIRECTOR',
     'Reviews weekly reports from all departments  →  Approves high-value POs  →  Makes strategic decisions based on financial statements.'),
]

for title, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(title + '  ')
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
    r2 = p.add_run(desc)
    r2.font.size = Pt(10); r2.font.color.rgb = DGRAY

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 16 — QUICK REFERENCE
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '16.  Common Tasks Quick Reference')

make_table(doc,
    ['Task', 'Responsible Role', 'Path in System'],
    [
        ['Add a new supplier',            'Purchasing Manager',       'Procurement → Suppliers → + Add Supplier'],
        ['Create a Purchase Order',       'Purchasing Manager',       'Procurement → Purchase Orders → + Create PO'],
        ['Approve a PO',                  'Purchasing Mgr / Accountant','Purchase Orders → click ✓ icon'],
        ['Inspect incoming goods',        'QA Officer',               'QA → Incoming Materials → + Add Record'],
        ['Create a GRN',                  'QA Officer',               'QA → Goods Receipt → + New GRN'],
        ['Check stock levels',            'Store Manager',            'Stores → Stock Inventory'],
        ['Transfer stock to van',         'Store Manager',            'Stores → Stock Transfers → + New Transfer'],
        ['Create a Job Order',            'Production Manager',       'Production → Job Orders → + New Job Order'],
        ['Request raw materials',         'Production Manager',       'Production → Material Requests → + New'],
        ['Record QA finished analysis',   'QA Officer',               'QA → Finished Products → + Add Analysis'],
        ['Release a QA batch',            'QA Officer',               'Finished Products → "Approve & Release"'],
        ['Add a new customer',            'Sales Manager',            'Sales → Customers → + Add Customer'],
        ['Register a van / route',        'Sales Manager',            'Sales → Routes & Vans → + New Van'],
        ['Create a sales invoice',        'Van Sales Rep',            'Sales → Invoices → + New Invoice'],
        ['Record a payment receipt',      'Van Sales Rep / Accountant','Sales → Receipts → + New Receipt'],
        ['Issue a credit note',           'Sales Manager',            'Sales → Credit Notes → + New Credit Note'],
        ['Create a dispatch record',      'Van Sales Rep',            'Sales → Dispatch Management → + New'],
        ['Post a journal entry',          'Accountant',               'Accounting → Journal Entries → + New'],
        ['Record an expense',             'Accountant',               'Accounting → Expenses → + New Expense'],
        ['View A/R Aging',                'Accountant',               'Accounting → A/R Aging'],
        ['Process payroll',               'Accountant + HR Manager',  'Accounting → Payroll / HR → Payroll Prep'],
        ['Record petty cash',             'Accountant',               'Accounting → Petty Cash → + New Entry'],
        ['Pay a supplier',                'Accountant',               'Accounting → Supplier Payments → + New'],
        ['Add an employee',               'HR Manager',               'HR → Employees → + Add Employee'],
        ['Approve leave',                 'HR Manager',               'HR → Leave Management → Review Request'],
        ['Submit weekly report',          'Department Head',          'Module page → "Send Weekly Report" button'],
        ['Review weekly reports',         'Managing Director',        'Weekly Reports → Review'],
        ['Add a system user',             'Super Admin',              'Settings → User Management → + Add User'],
        ['View audit trail',              'Super Admin',              'Settings → Audit Trail'],
        ['Export data to CSV',            'Any authorised user',      'Any module page → Export button'],
    ],
    col_widths=[2.1, 1.7, 2.7]
)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  IMPORTANT REMINDERS
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, 'Important Reminders for All Staff')

reminders = [
    ('Never share your login credentials.',
     'Each user has their own account. Sharing credentials violates company policy and compromises the audit trail.'),
    ('Always log out when you leave your workstation.',
     'Go to Settings → Profile → Sign Out, especially on shared computers.'),
    ('Issued invoices cannot be deleted.',
     'If a mistake is made, issue a Credit Note to reverse the transaction.'),
    ('Confirmed receipts cannot be deleted.',
     'If recorded in error, contact your Super Admin to void it.'),
    ('Stock deductions are automatic.',
     'When you save an invoice as Issued, stock is immediately deducted from the van. Ensure the van is properly loaded before issuing invoices.'),
    ('All actions are logged.',
     'Every create, edit, and delete is recorded in the Audit Trail with your name and timestamp. There is no hiding activity.'),
    ('Low stock alerts matter.',
     'When a product drops below its reorder level after a sale, the system shows a warning. Notify Stores and Purchasing immediately.'),
    ('Currency is Ghana Cedis (GH₵).',
     'All amounts in the system are in GH₵ unless otherwise specified.'),
    ('Batch numbers are critical.',
     'Always use the correct batch number when recording QA results. This is your traceability and recall record.'),
    ('Submit weekly reports without fail.',
     'The Managing Director relies on these for strategic decisions. Submit your department report every week.'),
]

for i, (title, detail) in enumerate(reminders, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f'{i}.  {title}  ')
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
    r2 = p.add_run(detail)
    r2.font.size = Pt(10); r2.font.color.rgb = DGRAY

# Footer note
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run('Aspee Pharmaceuticals ERP User Manual — Version 1.0  |  April 2026  |  Confidential — For Internal Use Only')
r_f.font.size = Pt(9)
r_f.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
r_f.italic = True

# ── Save ────────────────────────────────────────────────────────────────────
output_path = r'C:\Users\hp\Desktop\Developments\Aspee Pharmaceuticals\ASPEE_PHARMA_USER_MANUAL.docx'
doc.save(output_path)
print('Saved to: ' + output_path)
