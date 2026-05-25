"""
Generate individual department manuals for Aspee Pharmaceuticals ERP.
Produces one styled .docx per department in an 'Department Manuals' subfolder.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ──────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1e, 0x3a, 0x5f)
TEAL  = RGBColor(0x0f, 0x76, 0x6e)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DGRAY = RGBColor(0x33, 0x33, 0x33)
AMBER = RGBColor(0x92, 0x40, 0x0E)
SLATE = RGBColor(0x94, 0xA3, 0xB8)

# ── Style helpers ─────────────────────────────────────────────────────────────

def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(14); run.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), '1e3a5f')
    pBdr.append(bot); pPr.append(pBdr)

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(11); run.font.color.rgb = TEAL

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(10); run.font.color.rgb = DGRAY

def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text); run.font.size = Pt(10); run.font.color.rgb = DGRAY

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    run = p.add_run(text); run.font.size = Pt(10)

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent  = Inches(0.15); p.paragraph_format.right_indent = Inches(0.15)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'FEF3C7')
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run('  Note: ' + text)
    run.font.size = Pt(9.5); run.font.color.rgb = AMBER; run.italic = True

def warning(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent  = Inches(0.15); p.paragraph_format.right_indent = Inches(0.15)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'FEE2E2')
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run('  Important: ' + text)
    run.font.size = Pt(9.5); run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C); run.bold = True

def tip(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent  = Inches(0.15); p.paragraph_format.right_indent = Inches(0.15)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'DCFCE7')
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run('  Tip: ' + text)
    run.font.size = Pt(9.5); run.font.color.rgb = RGBColor(0x16, 0x4E, 0x32); run.italic = True

def table(doc, headers, rows, col_widths=None):
    n = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]; cell_bg(c, '1e3a5f')
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]
        pp.paragraph_format.space_before = Pt(3); pp.paragraph_format.space_after = Pt(3)
        run = pp.add_run(h); run.bold = True; run.font.size = Pt(9); run.font.color.rgb = WHITE
    for ri, row_data in enumerate(rows):
        r = tbl.rows[ri + 1]
        bg = 'F1F5F9' if ri % 2 == 0 else 'FFFFFF'
        for ci, txt in enumerate(row_data):
            c = r.cells[ci]; cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            pp = c.paragraphs[0]
            pp.paragraph_format.space_before = Pt(2); pp.paragraph_format.space_after = Pt(2)
            run = pp.add_run(str(txt)); run.font.size = Pt(9)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def pb(doc): doc.add_page_break()

def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width   = Inches(8.27); sec.page_height  = Inches(11.69)
    sec.left_margin  = Inches(1.0);  sec.right_margin = Inches(1.0)
    sec.top_margin   = Inches(1.0);  sec.bottom_margin= Inches(1.0)
    return doc

def cover_page(doc, dept_name, role_line, accent_hex='0f766e'):
    accent = RGBColor(int(accent_hex[0:2],16), int(accent_hex[2:4],16), int(accent_hex[4:6],16))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(55); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ASPEE PHARMACEUTICALS')
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = NAVY

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('ERP SYSTEM')
    r2.bold = True; r2.font.size = Pt(16); r2.font.color.rgb = DGRAY

    doc.add_paragraph()

    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(dept_name.upper() + ' USER MANUAL')
    r3.bold = True; r3.font.size = Pt(20); r3.font.color.rgb = accent

    doc.add_paragraph()

    for line, sz, bold in [
        ('Version 2.0  |  April 2026', 12, True),
        (role_line, 11, False),
        ('Currency: Ghana Cedis (GH₵)', 10, False),
        ('CONFIDENTIAL — FOR INTERNAL USE ONLY', 9, True),
    ]:
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pp.add_run(line)
        rr.font.size = Pt(sz); rr.bold = bold
        rr.font.color.rgb = AMBER if 'CONFIDENTIAL' in line else DGRAY

    pb(doc)

def section_getting_started(doc):
    h1(doc, '1.  Getting Started')
    h2(doc, 'Accessing the System')
    bullet(doc, 'Open Chrome or Edge and go to the ERP URL provided by your IT administrator.')
    bullet(doc, 'Enter your company email address and password, then click Sign In.')
    bullet(doc, 'If you forget your password, contact your Super Admin to reset it.')
    note(doc, 'If your account does not yet exist, contact the Super Admin to add you under Settings → User Management.')

    h2(doc, 'The Sidebar Navigation')
    bullet(doc, 'Click any section heading in the left sidebar to expand its sub-menu.')
    bullet(doc, 'The active page is highlighted in blue.')
    bullet(doc, 'The sidebar shows only the sections your role is permitted to access.')
    bullet(doc, 'Use the Collapse button at the bottom to hide labels and save screen space.')

    h2(doc, 'Common Page Actions')
    table(doc,
        ['Element', 'Purpose'],
        [
            ['Page Header',               'Title, subtitle, and breadcrumb trail'],
            ['Stat Cards',                'Key metrics at the top of each section'],
            ['Data Table',                'Searchable list of all records'],
            ['+ New / Create button',     'Opens the form to add a new record'],
            ['Eye icon',                  'Read-only detail view'],
            ['Pencil icon',               'Edit the record'],
            ['Trash icon',                'Delete the record (with confirmation)'],
            ['Export button',             'Download data as CSV'],
            ['Send Weekly Report button', 'Send a departmental summary to the MD'],
        ],
        col_widths=[2.2, 4.3]
    )
    pb(doc)

def section_reminders(doc, reminders):
    h1(doc, 'Important Reminders')
    for i, (title, detail) in enumerate(reminders, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f'{i}.  {title}  ')
        r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
        r2 = p.add_run(detail)
        r2.font.size = Pt(10); r2.font.color.rgb = DGRAY

def footer(doc, dept):
    doc.add_paragraph()
    pf = doc.add_paragraph(); pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run(
        f'Aspee Pharmaceuticals — {dept} User Manual  |  Version 2.0  |  April 2026  |  Confidential'
    )
    rf.font.size = Pt(9); rf.italic = True; rf.font.color.rgb = SLATE

OUT_DIR = r'C:\Users\hp\Desktop\Developments\Aspee Pharmaceuticals\Department Manuals'
os.makedirs(OUT_DIR, exist_ok=True)

# ════════════════════════════════════════════════════════════════════════════
# 1. PROCUREMENT DEPARTMENT
# ════════════════════════════════════════════════════════════════════════════
def build_procurement():
    doc = new_doc()
    cover_page(doc, 'Procurement Department', 'For: Purchasing Manager', '1e3a5f')
    section_getting_started(doc)

    h1(doc, '2.  Suppliers')
    body(doc, 'Path: Procurement → Suppliers')
    body(doc, 'The master record of all vendors Aspee Pharmaceuticals does business with.')
    h2(doc, 'To add a new supplier')
    bullet(doc, 'Click + Add Supplier.')
    bullet(doc, 'Fill in: Supplier Name, Contact Person, Phone, Email, Address, Payment Terms.')
    bullet(doc, 'Click Save.')
    table(doc,
        ['Field', 'Example Value'],
        [
            ['Supplier Name',  'Meridian Pharma Supplies, Tema'],
            ['Contact Person', 'Mr. Kwame Asante'],
            ['Phone',          '0244 123 456'],
            ['Email',          'orders@meridian.com'],
            ['Address',        'Tema Industrial Area, Block 5'],
            ['Payment Terms',  'Net 30 days'],
        ], col_widths=[2.3, 4.2])

    pb(doc)
    h1(doc, '3.  Purchase Orders (POs)')
    body(doc, 'Path: Procurement → Purchase Orders')
    body(doc, 'A Purchase Order is the formal document sent to a supplier to purchase goods.')
    h2(doc, 'Stat Cards')
    table(doc,
        ['Card', 'Meaning'],
        [
            ['Total POs',      'All purchase orders in the system'],
            ['Pending Review', 'POs awaiting approval'],
            ['Completed',      'Fully received POs'],
            ['Total Value',    'Combined GH₵ value of all POs'],
        ], col_widths=[2.0, 4.5])
    h2(doc, 'Status Lifecycle')
    body(doc, 'Pending  →  Approved  →  Shipped  →  Received  →  Closed')
    body(doc, '(Cancelled at any stage if rejected)')
    h2(doc, 'Approval Rules')
    bullet(doc, 'POs under GH₵ 10,000 — Approved by Purchasing Manager.')
    bullet(doc, 'POs over GH₵ 10,000  — Require Accountant (Finance) approval as well.')
    bullet(doc, 'Click the green ✓ icon on a Pending PO to approve it.')
    h2(doc, 'To create a PO')
    bullet(doc, 'Click + Create PO.')
    bullet(doc, 'Select the Supplier and enter the PO Number (e.g. PO-2026-001).')
    bullet(doc, 'Add line items: Product, Quantity, Unit Price, Unit of Measure.')
    bullet(doc, 'Review the total, then click Save.')
    h3(doc, 'Example Purchase Order')
    table(doc,
        ['Field', 'Value'],
        [
            ['PO Number',      'PO-2026-047'],
            ['Supplier',       'Meridian Pharma Supplies'],
            ['Date',           '10 April 2026'],
            ['Item 1',         'Paracetamol API — 500 kg @ GH₵ 120/kg = GH₵ 60,000'],
            ['Item 2',         'Maize Starch — 200 kg @ GH₵ 45/kg = GH₵ 9,000'],
            ['Total',          'GH₵ 69,000.00'],
            ['Approval Level', 'Finance (exceeds GH₵ 10,000)'],
        ], col_widths=[2.3, 4.2])

    pb(doc)
    h1(doc, '4.  Goods Receipt Note (GRN)')
    body(doc, 'Path: Quality Assurance → Goods Receipt  (also visible in QA sidebar)')
    body(doc, 'Created when a supplier delivers goods against a PO. Must be linked to the original PO. After QA approves the GRN items, stock is added to the warehouse.')
    h2(doc, 'GRN fields')
    bullet(doc, 'GRN Number — auto-generated reference')
    bullet(doc, 'Linked PO Number')
    bullet(doc, 'Supplier and delivery date')
    bullet(doc, 'Products received with quantity and batch number')
    bullet(doc, 'QA inspection result (Approved / Rejected / Quarantine)')

    h1(doc, '5.  Supplier Payments')
    body(doc, 'Path: Accounting → Supplier Payments  (also accessible from Procurement area)')
    body(doc, 'Record payments made to suppliers against outstanding Purchase Orders.')
    h3(doc, 'Example Payment')
    table(doc,
        ['Field', 'Value'],
        [
            ['Supplier',       'Meridian Pharma Supplies'],
            ['PO Reference',   'PO-2026-047'],
            ['Invoice Amount', 'GH₵ 69,000.00'],
            ['Amount Paid',    'GH₵ 34,500.00 (50% deposit)'],
            ['Balance',        'GH₵ 34,500.00'],
            ['Method',         'Bank Transfer'],
            ['Bank Reference', 'GCB-TRF-20260408'],
            ['Date Paid',      '08 April 2026'],
        ], col_widths=[2.3, 4.2])

    h1(doc, '6.  Requisitions (Internal Purchase Requests)')
    body(doc, 'Path: Procurement → Requisitions')
    body(doc, 'Stores or Production raise internal purchase requests when stock is needed. The Purchasing Manager reviews these and converts approved requests into formal Purchase Orders.')

    pb(doc)
    h1(doc, '7.  Quick Reference')
    table(doc,
        ['Task', 'Path'],
        [
            ['Add a supplier',          'Procurement → Suppliers → + Add Supplier'],
            ['Create a PO',             'Procurement → Purchase Orders → + Create PO'],
            ['Approve a PO',            'Purchase Orders → click ✓ icon on Pending row'],
            ['Record a GRN',            'QA → Goods Receipt → + New GRN'],
            ['Pay a supplier',          'Accounting → Supplier Payments → + New'],
            ['Review requisitions',     'Procurement → Requisitions'],
            ['Export PO data',          'Purchase Orders → Export button'],
            ['Send weekly report',      'Purchase Orders page → Send Weekly Report'],
        ], col_widths=[2.8, 3.7])

    section_reminders(doc, [
        ('Issued POs cannot be silently altered.', 'Any changes after approval must be documented.'),
        ('Approval tiers are enforced.', 'POs above GH₵ 10,000 require Finance co-approval. Do not bypass.'),
        ('GRN must match the PO.', 'Any shortfall in received quantity must be noted on the GRN.'),
        ('All actions are logged.', 'Every create, edit, and approve is recorded in the Audit Trail.'),
    ])
    footer(doc, 'Procurement Department')
    doc.save(os.path.join(OUT_DIR, '01_Procurement_Manual.docx'))
    print('Saved: 01_Procurement_Manual.docx')

# ════════════════════════════════════════════════════════════════════════════
# 2. QUALITY ASSURANCE
# ════════════════════════════════════════════════════════════════════════════
def build_qa():
    doc = new_doc()
    cover_page(doc, 'Quality Assurance', 'For: Quality Assurance Officer', '0f766e')
    section_getting_started(doc)

    h1(doc, '2.  QA Overview Dashboard')
    body(doc, 'Path: Quality Assurance → Overview')
    body(doc, 'Shows the live status of all QA activities — incoming inspections, in-process checks, and finished product releases.')

    h1(doc, '3.  Incoming Materials Inspection')
    body(doc, 'Path: Quality Assurance → Incoming Materials')
    body(doc, 'Every supplier delivery must be inspected before materials are accepted into the warehouse. The outcome determines whether goods enter stock.')
    h2(doc, 'Process')
    bullet(doc, 'Procurement creates a GRN when goods arrive.')
    bullet(doc, 'QA Officer inspects the delivery against the PO and GRN.')
    bullet(doc, 'Record the inspection result: Approved, Rejected, or Quarantine.')
    bullet(doc, 'Approved materials are automatically added to warehouse stock.')
    h2(doc, 'Fields to complete')
    table(doc,
        ['Field', 'Description'],
        [
            ['GRN Number',    'Reference to the linked Goods Receipt Note'],
            ['Supplier',      'Name of the delivering supplier'],
            ['Product',       'Item being inspected'],
            ['Batch Number',  'Manufacturer\'s batch code on the packaging'],
            ['Qty Received',  'Actual quantity received (may differ from ordered)'],
            ['Test Results',  'Appearance, assay, moisture, purity as applicable'],
            ['Final Status',  'Approved / Rejected / Quarantine'],
        ], col_widths=[2.0, 4.5])
    h3(doc, 'Example Incoming Record')
    table(doc,
        ['Field', 'Value'],
        [
            ['GRN Number',   'GRN-2026-022'],
            ['Supplier',     'Meridian Pharma Supplies'],
            ['Product',      'Paracetamol API'],
            ['Batch Number', 'BATCH-PAR-0041'],
            ['Qty Ordered',  '500 kg'],
            ['Qty Received', '498 kg'],
            ['Appearance',   'White crystalline powder — Pass'],
            ['Assay',        '99.2% — Pass'],
            ['QA Status',    'Approved'],
        ], col_widths=[2.3, 4.2])

    pb(doc)
    h1(doc, '4.  Goods Receipt Note (GRN)')
    body(doc, 'Path: Quality Assurance → Goods Receipt')
    body(doc, 'The official record that goods have been received against a PO. The QA team creates and signs off GRNs. Once approved, stock is added automatically.')
    note(doc, 'GRN is accessible from both the QA sidebar and the Procurement area, since both teams interact with it.')

    h1(doc, '5.  In-Process Controls')
    body(doc, 'Path: Quality Assurance → In Process Controls')
    body(doc, 'QA performs checks at defined stages during manufacturing to confirm the product is being made correctly.')
    h3(doc, 'Example — Tablet Compression Check')
    table(doc,
        ['Check Point', 'Parameter', 'Acceptable Range', 'Example Result'],
        [
            ['Weight',         'Average tablet weight', '500 mg ± 5%',  '502 mg — Pass'],
            ['Hardness',       'Tablet hardness',       '4–8 kp',       '6.2 kp — Pass'],
            ['Disintegration', 'Disintegration time',   '< 15 minutes',      '8 min — Pass'],
            ['Appearance',     'Visual inspection',     'No defects',        'Pass'],
        ], col_widths=[1.6, 1.8, 1.7, 1.4])

    pb(doc)
    h1(doc, '6.  Finished Products Analysis')
    body(doc, 'Path: Quality Assurance → Finished Products')
    body(doc, 'Every manufactured batch must pass QA final analysis before it can be released for sale.')
    h2(doc, 'Outcome options')
    table(doc,
        ['Status', 'Meaning', 'Next Step'],
        [
            ['Passed',    'All tests within specification',              'Set Release Date — batch enters Finished Goods stock'],
            ['Failed',    'Tests outside specification',                  'Quarantine batch — notify Production Manager'],
            ['Quarantine','Under investigation — hold for decision', 'Investigate root cause, then re-evaluate'],
        ], col_widths=[1.4, 2.5, 2.6])
    h2(doc, 'To add a Finished Products analysis record')
    bullet(doc, 'Click + Add Analysis Record.')
    bullet(doc, 'Select the Production Order (Job Order) this batch belongs to.')
    bullet(doc, 'Enter Product Name, Batch Number, Analyst Name, and Analysis Date.')
    bullet(doc, 'Record all test results (dissolution, assay, microbial, etc.).')
    bullet(doc, 'Set Overall Status (Passed / Failed / Quarantine).')
    bullet(doc, 'If Passed, set the Release Date.')
    bullet(doc, 'Click Save.')
    tip(doc, 'For Quarantined batches that later clear review, use the green "Approve & Release" button directly from the table — no need to re-open the full record.')
    h3(doc, 'Example Finished Product Record')
    table(doc,
        ['Field', 'Value'],
        [
            ['Product',       'Paracetamol 500mg Tablets'],
            ['Batch Number',  'BPC-2026-018'],
            ['Job Order',     'JO-2026-011'],
            ['Analyst',       'Abena Mensah'],
            ['Analysis Date', '05 April 2026'],
            ['Dissolution',   '85% at 30 min — Pass'],
            ['Assay',         '98.8% — Pass'],
            ['Microbial',     'TAMC < 1,000 cfu/g — Pass'],
            ['Status',        'Passed'],
            ['Release Date',  '07 April 2026'],
        ], col_widths=[2.3, 4.2])

    pb(doc)
    h1(doc, '7.  Internal QA Reports')
    body(doc, 'Path: Quality Assurance → Internal Reports')
    body(doc, 'Store deviation reports, non-conformance records, and corrective action summaries here.')

    h1(doc, '8.  Compliance Monitoring')
    body(doc, 'Path: Compliance → Regulators & Renewals')
    body(doc, 'Track all regulatory licences and renewal dates. QA Officers can view and update compliance records.')
    table(doc,
        ['Regulator', 'Licence', 'Expiry', 'Status'],
        [
            ['FDA Ghana', "Manufacturer's Licence", '31 Dec 2026', 'Active'],
            ['EPA',       'Factory Registration',   '30 Jun 2026', 'Due for Renewal'],
        ], col_widths=[1.5, 2.3, 1.5, 1.2])
    note(doc, 'Set calendar reminders 90 days before any licence expiry to initiate renewal.')

    pb(doc)
    h1(doc, '9.  Quick Reference')
    table(doc,
        ['Task', 'Path'],
        [
            ['Inspect incoming goods',          'QA → Incoming Materials → + Add Record'],
            ['Create a GRN',                    'QA → Goods Receipt → + New GRN'],
            ['Record in-process check',         'QA → In Process Controls → + Add Record'],
            ['Add finished products analysis',  'QA → Finished Products → + Add Analysis'],
            ['Release a passed batch',          'Finished Products → "Approve & Release" button'],
            ['Log a deviation report',          'QA → Internal Reports → + New Report'],
            ['View compliance licences',        'Compliance → Regulators & Renewals'],
            ['Send weekly report',              'QA Overview page → Send Weekly Report'],
        ], col_widths=[2.8, 3.7])

    section_reminders(doc, [
        ('Batch numbers are mandatory.', 'Every inspection, analysis, and finished-product record must carry the correct batch number — this is your traceability chain.'),
        ('No release without QA sign-off.', 'Finished goods may not enter saleable stock until a QA Analyst records a Passed status and sets a Release Date.'),
        ('Contamination and critical defects must be escalated.', 'Raise a Non-Conformance in Internal Audit and notify the QA Manager immediately.'),
        ('All actions are logged.', 'Every record you create or edit is timestamped with your name in the Audit Trail.'),
    ])
    footer(doc, 'Quality Assurance')
    doc.save(os.path.join(OUT_DIR, '02_QA_Manual.docx'))
    print('Saved: 02_QA_Manual.docx')

# ════════════════════════════════════════════════════════════════════════════
# 3. STORES DEPARTMENT
# ════════════════════════════════════════════════════════════════════════════
def build_stores():
    doc = new_doc()
    cover_page(doc, 'Stores Department', 'For: Store Manager', '7c3aed')
    section_getting_started(doc)

    h1(doc, '2.  Products (Product Master)')
    body(doc, 'Path: Stores → Products')
    body(doc, 'The master list of every product or material the company handles. All items must exist here before they can be stocked, transferred, or sold.')
    table(doc,
        ['Field', 'Description', 'Example'],
        [
            ['Product Name',  'Full item name',                     'Paracetamol 500mg Tablets'],
            ['SKU',           'Unique stock-keeping code',           'PAR-500-TAB'],
            ['Unit',          'Unit of measurement',                 'Bottles / Kg / Boxes'],
            ['Material Type', 'Raw Material / Finished / Packaging', 'Finished Product'],
            ['Reorder Level', 'Minimum qty before alert fires',      '500'],
        ], col_widths=[1.7, 2.5, 2.3])
    tip(doc, 'Always set a realistic reorder level. When stock drops to or below this number, the system automatically alerts you and Purchasing.')

    pb(doc)
    h1(doc, '3.  Stock Inventory')
    body(doc, 'Path: Stores → Stock Inventory')
    body(doc, 'Real-time view of stock across all locations: warehouse, production floor, sales vans.')
    table(doc,
        ['Status', 'Meaning', 'Action'],
        [
            ['Adequate',     'Above reorder level',           'Monitor normally'],
            ['Warning',      '1x – 1.5x reorder level',        'Plan a purchase request'],
            ['Low',          'At or below reorder level',     'Raise a Purchase Request immediately'],
            ['Out of Stock', 'Zero quantity',                 'Emergency Purchase Request'],
        ], col_widths=[1.4, 2.3, 2.8])
    note(doc, 'Use the material-type tab buttons (All / Raw Material / Finished Product / Packaging) to filter the view.')

    h1(doc, '4.  Internal Use')
    body(doc, 'Path: Stores → Internal Use  —  NEW in v2.0')
    body(doc, 'Records stock consumed internally — not sold, but used by the company. Saving a record immediately deducts the quantity from stock.')
    table(doc,
        ['Purpose', 'Example Use'],
        [
            ['Office Use',               'Packaging materials used by admin'],
            ['Testing / Sampling',       'Products used for QA samples'],
            ['Staff Consumption',        'Medicines issued to factory floor workers'],
            ['Company Promotion',        'Products given as promotional samples'],
            ['Research & Development',   'Materials consumed in R&D'],
            ['Other',                    'Any other internal consumption'],
        ], col_widths=[2.3, 4.2])
    warning(doc, 'Deleting an Internal Use record does NOT restore the stock. Always verify quantities before saving.')

    pb(doc)
    h1(doc, '5.  Material Defects')
    body(doc, 'Path: Stores → Material Defects  —  NEW in v2.0')
    body(doc, 'Records defective, damaged, or contaminated materials. Stock is automatically written off when the record is saved.')
    table(doc,
        ['Defect Type', 'Severity'],
        [
            ['Breakage',             'Warning'],
            ['Leakage',              'Warning'],
            ['Spillage',             'Danger'],
            ['Contamination',        'Critical — escalate to QA immediately'],
            ['Damaged Packaging',    'Warning'],
            ['Manufacturing Defect', 'Critical — escalate to QA immediately'],
        ], col_widths=[2.5, 4.0])
    warning(doc, 'Critical defects (Contamination / Manufacturing Defect) must be reported to QA and a Non-Conformance raised in Internal Audit.')

    h1(doc, '6.  Material Expiry')
    body(doc, 'Path: Stores → Material Expiry  —  NEW in v2.0')
    body(doc, 'Records materials that have reached or passed their expiry date. Stock is automatically removed when the record is saved.')
    table(doc,
        ['Disposal Method', 'Description'],
        [
            ['Destroyed / Incinerated',           'Completely destroyed on-site'],
            ['Returned to Supplier',              'Sent back to the original supplier'],
            ['Quarantined',                       'Held pending regulatory decision'],
            ['Donated',                           'Donated to an approved charitable body'],
            ['Disposed via Licensed Waste Handler','Handled by a certified waste contractor'],
        ], col_widths=[2.5, 4.0])
    note(doc, 'Always record the batch number and expiry date shown on the packaging — this is a regulatory traceability requirement.')

    pb(doc)
    h1(doc, '7.  Stock Transfers')
    body(doc, 'Path: Stores → Stock Transfers')
    body(doc, 'Move stock between locations — e.g. Warehouse → Production Floor, or Warehouse → Sales Department store.')
    h2(doc, 'To create a transfer')
    bullet(doc, 'Click + New Transfer.')
    bullet(doc, 'Select From Location and To Location.')
    bullet(doc, 'Add the products and quantities to transfer.')
    bullet(doc, 'Click Save.')

    h1(doc, '8.  Material Requests')
    body(doc, 'Path: Stores → Material Requests')
    body(doc, 'Production raises Material Requests when they need raw materials for a Job Order. Review the request, check stock, and issue materials if available. If not, raise a Purchase Request.')

    h1(doc, '9.  Purchase Requests')
    body(doc, 'Path: Stores → Purchase Requests')
    body(doc, 'Raise a Purchase Request when a needed material is out of stock. This notifies the Purchasing Manager to create a formal Purchase Order.')

    h1(doc, '10.  Sales Movements (Read-Only)')
    body(doc, 'Path: Stores → Sales Movements')
    body(doc, 'A read-only view showing which products left which van location, when, and against which invoice. Use this to reconcile van stock.')

    h1(doc, '11.  QA Reports (Stores View)')
    body(doc, 'Path: Stores → QA Reports')
    body(doc, 'View all QA inspection results for received materials and released finished products without needing full QA module access.')

    pb(doc)
    h1(doc, '12.  Quick Reference')
    table(doc,
        ['Task', 'Path'],
        [
            ['Add a product',              'Stores → Products → + Add Product'],
            ['Check stock levels',         'Stores → Stock Inventory'],
            ['Log internal use',           'Stores → Internal Use → + Log Internal Use'],
            ['Log a material defect',      'Stores → Material Defects → + Log Defect'],
            ['Log expired materials',      'Stores → Material Expiry → + Log Expiry'],
            ['Transfer stock',             'Stores → Stock Transfers → + New Transfer'],
            ['Review material requests',   'Stores → Material Requests'],
            ['Raise a purchase request',   'Stores → Purchase Requests → + New'],
            ['View sales movements',       'Stores → Sales Movements'],
            ['View QA results',            'Stores → QA Reports'],
            ['Send weekly report',         'Stock Inventory page → Send Weekly Report'],
        ], col_widths=[2.8, 3.7])

    section_reminders(doc, [
        ('Internal Use, Defects, and Expiry records are one-way.', 'Stock is deducted immediately on save. Deleting the record does NOT restore stock. Double-check quantities before saving.'),
        ('Reorder levels must be accurate.', 'If a product has no reorder level set, the system cannot alert you. Keep product master data up to date.'),
        ('GRN approval is the gateway.', 'Materials cannot enter stock without a QA-approved GRN. Do not accept stock informally.'),
        ('All actions are logged.', 'Every transfer, issue, and adjustment is recorded in the Audit Trail with your name and timestamp.'),
    ])
    footer(doc, 'Stores Department')
    doc.save(os.path.join(OUT_DIR, '03_Stores_Manual.docx'))
    print('Saved: 03_Stores_Manual.docx')

# ════════════════════════════════════════════════════════════════════════════
# 4. PRODUCTION DEPARTMENT
# ════════════════════════════════════════════════════════════════════════════
def build_production():
    doc = new_doc()
    cover_page(doc, 'Production Department', 'For: Production Manager', 'b45309')
    section_getting_started(doc)

    h1(doc, '2.  Bill of Materials (BOM)')
    body(doc, 'Path: Production → Bill of Materials')
    body(doc, 'The BOM is the recipe for producing a finished product — which raw materials are needed, in what quantities, and in what units.')
    h2(doc, 'To create a BOM')
    bullet(doc, 'Click + New BOM.')
    bullet(doc, 'Select the Finished Product (e.g. Paracetamol 500mg Tablets).')
    bullet(doc, 'Add each raw material: name, quantity per batch, unit of measure.')
    bullet(doc, 'Save.')
    h3(doc, 'Example BOM — Paracetamol 500mg Tablets (10,000 tablet batch)')
    table(doc,
        ['Component', 'Quantity', 'Unit'],
        [
            ['Paracetamol API',              '5.00', 'kg'],
            ['Maize Starch',                 '1.20', 'kg'],
            ['Microcrystalline Cellulose',   '0.80', 'kg'],
            ['Magnesium Stearate',           '0.05', 'kg'],
            ['Talc',                         '0.10', 'kg'],
            ['PVC Blister Foil',             '200',  'sheets'],
        ], col_widths=[3.5, 1.5, 1.5])

    pb(doc)
    h1(doc, '3.  Job Orders (Production Orders)')
    body(doc, 'Path: Production → Job Orders')
    body(doc, 'A Job Order is the formal instruction to manufacture a specific batch of a product.')
    h2(doc, 'Status lifecycle')
    body(doc, 'Planned  →  In Progress  →  Completed  →  Released (after QA Approval)')
    h2(doc, 'To create a Job Order')
    bullet(doc, 'Click + New Job Order.')
    bullet(doc, 'Enter the Order Number (e.g. JO-2026-015).')
    bullet(doc, 'Select the Product to manufacture.')
    bullet(doc, 'Enter the Batch Size (quantity to produce).')
    bullet(doc, 'Set the Planned Start Date and Expected Completion Date.')
    bullet(doc, 'The system automatically shows the BOM for that product.')
    bullet(doc, 'Save.')
    h3(doc, 'Example Job Order')
    table(doc,
        ['Field', 'Value'],
        [
            ['Order Number', 'JO-2026-015'],
            ['Product',      'Paracetamol 500mg Tablets'],
            ['Batch Size',   '50,000 tablets'],
            ['Start Date',   '08 April 2026'],
            ['Expected End', '12 April 2026'],
            ['Status',       'In Progress'],
        ], col_widths=[2.3, 4.2])

    h1(doc, '4.  Production Material Requests')
    body(doc, 'Path: Production → Material Requests')
    body(doc, 'When starting a Job Order, raise a Material Request to get raw materials from Stores.')
    h2(doc, 'Process')
    bullet(doc, '1. Production Manager creates a Material Request linked to the Job Order.')
    bullet(doc, '2. Store Manager reviews the request and checks stock availability.')
    bullet(doc, '3. If available, Stores issues the materials — stock quantity decreases.')
    bullet(doc, '4. If not available, Store Manager raises a Purchase Request to Procurement.')
    bullet(doc, '5. Production uses the materials to manufacture the batch.')
    bullet(doc, '6. Completed batch goes to QA for Finished Products Analysis.')
    note(doc, 'Production cannot begin until Stores confirms the Material Request has been issued. Check with the Store Manager before starting the manufacturing run.')

    pb(doc)
    h1(doc, '5.  Quick Reference')
    table(doc,
        ['Task', 'Path'],
        [
            ['View/edit BOMs',              'Production → Bill of Materials'],
            ['Create a new BOM',            'Production → Bill of Materials → + New BOM'],
            ['Create a Job Order',          'Production → Job Orders → + New Job Order'],
            ['Update Job Order status',     'Production → Job Orders → Edit record'],
            ['Request raw materials',       'Production → Material Requests → + New Request'],
            ['Check stock availability',   'Stores → Stock Inventory (Store Manager access)'],
            ['Send weekly report',          'Production → Job Orders page → Send Weekly Report'],
        ], col_widths=[2.8, 3.7])

    section_reminders(doc, [
        ('Always link material requests to the Job Order.', 'Unlinked requests may be delayed or rejected by Stores.'),
        ('QA release is mandatory.', 'A completed batch cannot be sold until QA records a Passed status and sets a release date.'),
        ('Batch numbers matter.', 'The batch number used in QA records must match what was manufactured. Agree the batch number with QA before production starts.'),
        ('All actions are logged.', 'Every Job Order update and material request is tracked with your name and timestamp.'),
    ])
    footer(doc, 'Production Department')
    doc.save(os.path.join(OUT_DIR, '04_Production_Manual.docx'))
    print('Saved: 04_Production_Manual.docx')

# ════════════════════════════════════════════════════════════════════════════
# 5. SALES DEPARTMENT
# ════════════════════════════════════════════════════════════════════════════
def build_sales():
    doc = new_doc()
    cover_page(doc, 'Sales Department', 'For: Sales Manager & Van Sales Representatives', '16a34a')
    section_getting_started(doc)

    h1(doc, '2.  Customers')
    body(doc, 'Path: Sales → Customers')
    body(doc, 'All customers must be registered here before an invoice can be raised. The ERP links every invoice to a registered customer.')
    h2(doc, 'Customer fields')
    table(doc,
        ['Field', 'Description'],
        [
            ['Customer Name',    'Full business name'],
            ['Contact Person',   'Primary contact at the customer'],
            ['Phone',            'Contact number'],
            ['Customer Category','Wholesale / Retail / Hospital / Pharmacy / Institution'],
            ['Customer Location','Town or area (e.g. Kumasi, Ashanti)'],
            ['Sales Person',     'The sales rep assigned to this customer'],
            ['Route',            'Delivery route this customer belongs to'],
            ['Credit Limit',     'Maximum credit allowed (GH₵)'],
            ['Payment Terms',    'Net 7 / Net 14 / Net 30 days'],
        ], col_widths=[2.3, 4.2])
    tip(doc, 'Use the Import button to bulk-upload customers from an Excel file. Download the template from the import screen first.')
    tip(doc, 'Click the SOA icon on any customer row to download their full Statement of Account.')

    pb(doc)
    h1(doc, '3.  Routes & Vans')
    body(doc, 'Path: Sales → Routes & Vans')
    body(doc, 'Each Van is assigned to a delivery route/territory. The van\'s route determines which stock location is deducted when invoices are issued.')
    table(doc,
        ['Field', 'Example'],
        [
            ['Van ID',       'VAN-001'],
            ['Driver',       'Kwaku Asante'],
            ['Plate Number', 'GR 1234-23'],
            ['Route Area',   'Accra Central'],
        ], col_widths=[2.3, 4.2])

    h1(doc, '4.  Stock Movements (Van Loading)')
    body(doc, 'Path: Sales → Stock Movements  —  NEW in v2.0')
    body(doc, 'Before a Van Sales Rep can issue invoices, stock must be loaded from the Sales Department store onto their van.')
    h2(doc, 'To load a van')
    bullet(doc, 'Click + Load Van.')
    bullet(doc, 'From Location: select the Sales Department store.')
    bullet(doc, 'To Location: select the specific van (e.g. Sales Van — VAN-001).')
    bullet(doc, 'Add products and quantities.')
    bullet(doc, 'Click Save.')
    warning(doc, 'The system only allows transfers FROM the Sales Department store TO an individual van. Reverse flows are blocked. Contact Stores for other transfer types.')
    h2(doc, 'Movement types shown in the table')
    table(doc,
        ['Type', 'What it shows'],
        [
            ['Van Load',     'Stock moved from Sales Department store to a van'],
            ['Sales Invoice','Stock deducted from a van when an invoice was issued to a customer'],
        ], col_widths=[2.0, 4.5])

    pb(doc)
    h1(doc, '5.  Sales Invoices')
    body(doc, 'Path: Sales → Invoices')
    body(doc, 'An invoice represents a sale made to a customer. The system deducts stock from the van automatically when an invoice is saved as Issued.')
    h2(doc, 'Invoice types')
    table(doc,
        ['Type', 'Description'],
        [
            ['Cash Sale',   'Payment received immediately at point of sale.'],
            ['Credit Sale', 'Payment due on a future date based on credit terms.'],
        ], col_widths=[2.0, 4.5])
    h2(doc, 'Status lifecycle')
    body(doc, 'Draft  →  Issued  →  Partially Paid  →  Paid')
    body(doc, '(Overdue if payment due date passes)  |  (Cancelled only via Credit Note)')
    h2(doc, 'To create an invoice')
    bullet(doc, '1. Click + New Invoice.')
    bullet(doc, '2. Select the Customer from the searchable dropdown.')
    bullet(doc, '   → The system auto-fills: Location, Route, Salesperson, Credit Limit, and Outstanding Balance.')
    bullet(doc, '3. Select the Route / Van (determines which van\'s stock is deducted).')
    bullet(doc, '4. Enter Invoice Date and Due Date.')
    bullet(doc, '5. Select Invoice Type (Cash Sale or Credit Sale).')
    bullet(doc, '6. Add line items:')
    bullet(doc, '   → Product, Quantity, Unit Price', level=1)
    bullet(doc, '   → Discount Amount per line (optional)', level=1)
    bullet(doc, '   → Batch Number (required for controlled drugs)', level=1)
    bullet(doc, '   → Tick Damaged if the product is defective (excluded from revenue)', level=1)
    bullet(doc, '   → Tick Gifted if given free of charge (excluded from revenue)', level=1)
    bullet(doc, '7. Review totals (subtotal, discount, grand total).')
    bullet(doc, '8. Set status to Issued when ready (or Draft to save without committing).')
    bullet(doc, '9. Click Save.')
    note(doc, 'Only Draft invoices can be deleted. Once Issued, use a Credit Note to reverse.')
    warning(doc, 'When an invoice is saved as Issued: stock is immediately deducted from the van, a journal entry is posted, and a low-stock alert fires if any product falls below its reorder level.')

    pb(doc)
    h1(doc, '6.  Waybills')
    body(doc, 'Path: Sales → Waybills  —  NEW in v2.0')
    body(doc, 'A Waybill is a document accompanying goods loaded onto a van for a route run. Generate one each time you load a van.')
    bullet(doc, 'Click + Generate Waybill.')
    bullet(doc, 'Select the Van and confirm driver details.')
    bullet(doc, 'Enter the Date and Sales Person Name.')
    bullet(doc, 'Save — the system calculates the total load value.')

    h1(doc, '7.  Dispatch Management')
    body(doc, 'Path: Sales → Dispatch Management')
    body(doc, 'Dispatch records track the physical delivery of goods to customers.')
    body(doc, 'Status: Draft  →  Pending  →  In Transit  →  Completed  (or Cancelled)')
    bullet(doc, 'Enter Dispatch Number, Van, Driver, dispatch date.')
    bullet(doc, 'Link the invoices being dispatched.')
    bullet(doc, 'Update to In Transit when goods leave. Mark Completed on confirmed delivery.')

    h1(doc, '8.  Receipts (Payment Collection)')
    body(doc, 'Path: Sales → Receipts')
    body(doc, 'Record payments collected from customers against their invoices.')
    table(doc,
        ['Payment Method', 'Reference Required'],
        [
            ['Cash',          'None required'],
            ['Cheque',        'Cheque number and bank'],
            ['Mobile Money',  'MoMo transaction reference'],
            ['Bank Transfer', 'Bank reference / SWIFT number'],
        ], col_widths=[2.3, 4.2])
    warning(doc, 'Confirmed or Cleared receipts cannot be deleted. Contact your Super Admin to void if recorded in error.')

    pb(doc)
    h1(doc, '9.  Credit Notes')
    body(doc, 'Path: Sales → Credit Notes')
    body(doc, 'Issue a Credit Note when a customer returns goods or an issued invoice must be partially or fully reversed.')
    bullet(doc, 'Click + New Credit Note.')
    bullet(doc, 'Link to the original Invoice.')
    bullet(doc, 'Enter the reason, items being credited, and amounts.')
    bullet(doc, 'Save.')

    h1(doc, '10.  Sales Reports (11 Reports)')
    body(doc, 'Path: Sales → Sales Reports')
    table(doc,
        ['Report Tab', 'What it shows'],
        [
            ['Distribution',        'All invoices by customer — volume, revenue, route, category'],
            ['Stock by Salesperson', 'Current van stock on hand per sales rep'],
            ['Stock by Route',       'Current van stock on hand per delivery route'],
            ['Debtors by Staff',     'Outstanding amounts owed, grouped by salesperson'],
            ['Debtors by Route',     'Outstanding amounts owed, grouped by route'],
            ['Debtors by Period',    'Aging buckets: 0–30, 31–60, 61–90, 90+ days'],
            ['Sales vs Cash',        'Total sales value vs cash collected — credit exposure'],
            ['Cheques Received',     'All cheque payments with cheque number and bank details'],
            ['Shortage / Excess',    'Stock given to van vs stock sold — identifies missing stock'],
            ['Requisitions',         'Stock requests raised by sales reps from Stores'],
            ['Product Distribution', 'Which products are moving to which customers and routes'],
        ], col_widths=[2.3, 4.2])

    pb(doc)
    h1(doc, '11.  Quick Reference')
    table(doc,
        ['Task', 'Who', 'Path'],
        [
            ['Add a customer',              'Sales Manager',  'Sales → Customers → + Add Customer'],
            ['Import customers (Excel)',     'Sales Manager',  'Sales → Customers → Import'],
            ['Export customer SOA',         'Sales Manager',  'Sales → Customers → SOA icon on row'],
            ['Register a van / route',      'Sales Manager',  'Sales → Routes & Vans → + New Van'],
            ['Load stock onto a van',       'Sales Manager',  'Sales → Stock Movements → + Load Van'],
            ['Generate a waybill',          'Sales Manager',  'Sales → Waybills → + Generate Waybill'],
            ['Create a sales invoice',      'Van Sales Rep',  'Sales → Invoices → + New Invoice'],
            ['Record a receipt',            'Van Sales Rep',  'Sales → Receipts → + New Receipt'],
            ['Issue a credit note',         'Sales Manager',  'Sales → Credit Notes → + New Credit Note'],
            ['Create a dispatch record',    'Van Sales Rep',  'Sales → Dispatch Management → + New'],
            ['View all 11 sales reports',   'Sales Manager',  'Sales → Sales Reports'],
            ['Send weekly report',          'Sales Manager',  'Sales → Invoices page → Send Weekly Report'],
        ], col_widths=[2.3, 1.5, 2.7])

    section_reminders(doc, [
        ('Load the van before issuing invoices.', 'A van with zero stock will block invoice creation. Always confirm with Stores that the van has been loaded via Sales → Stock Movements.'),
        ('Issued invoices cannot be deleted.', 'If a mistake is made after issuing, raise a Credit Note to reverse the transaction.'),
        ('Damaged and Gifted items are still stocked-out.', 'Marking a line as Damaged or Gifted excludes it from revenue, but stock is still deducted from the van.'),
        ('Confirmed receipts cannot be deleted.', 'Contact the Super Admin to void any receipt recorded in error.'),
        ('Batch numbers are required for controlled drugs.', 'Always enter the correct batch number on each invoice line for controlled drug items.'),
        ('Submit your weekly report every week.', 'The Managing Director reviews all department reports. Missing submissions will be escalated.'),
    ])
    footer(doc, 'Sales Department')
    doc.save(os.path.join(OUT_DIR, '05_Sales_Manual.docx'))
    print('Saved: 05_Sales_Manual.docx')

# ════════════════════════════════════════════════════════════════════════════
# 6. ACCOUNTING DEPARTMENT
# ════════════════════════════════════════════════════════════════════════════
def build_accounting():
    doc = new_doc()
    cover_page(doc, 'Accounting Department', 'For: Accountant', '0369a1')
    section_getting_started(doc)

    h1(doc, '2.  Journal Entries')
    body(doc, 'Path: Accounting → Journal Entries')
    body(doc, 'Manual journal entries for transactions not automatically posted by the system. Most invoice and receipt transactions post automatically.')
    table(doc,
        ['Account', 'Type', 'Normal Balance'],
        [
            ['Sales Revenue',      'Income',    'Credit'],
            ['Accounts Receivable','Asset',     'Debit'],
            ['Cash & Bank',        'Asset',     'Debit'],
            ['Inventory / Stock',  'Asset',     'Debit'],
            ['Cost of Goods Sold', 'Expense',   'Debit'],
            ['Accounts Payable',   'Liability', 'Credit'],
            ['VAT Payable',        'Liability', 'Credit'],
        ], col_widths=[2.3, 1.5, 2.7])

    h1(doc, '3.  General Ledger')
    body(doc, 'Path: Accounting → General Ledger')
    body(doc, 'All posted journal entries organised by account. Use this to trace any transaction back to its origin.')

    h1(doc, '4.  Trial Balance')
    body(doc, 'Path: Accounting → Trial Balance')
    body(doc, 'Summary of all account balances confirming total debits equal total credits. Run at month-end to verify book integrity before producing financial statements.')

    pb(doc)
    h1(doc, '5.  Expenses')
    body(doc, 'Path: Accounting → Expenses')
    body(doc, 'Record and categorise all company operational expenses.')
    table(doc,
        ['Field', 'Example'],
        [
            ['Category',       'Fuel & Transport'],
            ['Description',    'Fuel for Van 1 delivery route — April 2026'],
            ['Amount',         'GH₵ 850.00'],
            ['Date',           '10 April 2026'],
            ['Payment Method', 'Cash (Petty Cash)'],
        ], col_widths=[2.3, 4.2])

    h1(doc, '6.  Payroll')
    body(doc, 'Path: Accounting → Payroll')
    body(doc, 'Monthly payroll processing. Works with HR\'s payroll preparation data.')
    table(doc,
        ['Employee', 'Basic (GH₵)', 'Transport', 'SSNIT 5.5%', 'Tax', 'Net Pay'],
        [
            ['Abena Mensah (QA)', '3,800', '400', '209.00', '312.00', '3,679.00'],
            ['Kwaku Asante (Van)', '2,500', '300', '137.50', '178.00', '2,484.50'],
        ], col_widths=[1.9, 1.0, 1.0, 1.1, 0.9, 1.0])

    h1(doc, '7.  Tax Periods')
    body(doc, 'Path: Accounting → Tax Periods')
    body(doc, 'Track VAT filing periods, corporate tax assessments, and PAYE submissions to GRA.')

    h1(doc, '8.  Petty Cash')
    body(doc, 'Path: Accounting → Petty Cash')
    body(doc, 'Track small cash disbursements for day-to-day operational expenses.')
    table(doc,
        ['Date', 'Description', 'Amount (GH₵)', 'Balance (GH₵)'],
        [
            ['07 Apr 2026', 'Office stationery', '85.00', '415.00'],
            ['08 Apr 2026', 'Factory floor refreshments', '45.00', '370.00'],
            ['10 Apr 2026', 'Courier delivery fee', '30.00', '340.00'],
        ], col_widths=[1.4, 2.5, 1.5, 1.5])

    pb(doc)
    h1(doc, '9.  Supplier Payments')
    body(doc, 'Path: Accounting → Supplier Payments')
    body(doc, 'Record all payments made to suppliers against outstanding Purchase Orders.')
    table(doc,
        ['Field', 'Example'],
        [
            ['Supplier',       'Meridian Pharma Supplies'],
            ['PO Reference',   'PO-2026-047'],
            ['Invoice Amount', 'GH₵ 69,000.00'],
            ['Amount Paid',    'GH₵ 34,500.00'],
            ['Balance',        'GH₵ 34,500.00'],
            ['Method',         'Bank Transfer'],
            ['Bank Reference', 'GCB-TRF-20260408'],
        ], col_widths=[2.3, 4.2])

    h1(doc, '10.  Collections')
    body(doc, 'Path: Accounting → Collections')
    body(doc, 'The accounting view of customer payments. Cross-references with Sales Receipts to confirm banking and reconcile cash received.')

    h1(doc, '11.  A/R Aging Report')
    body(doc, 'Path: Accounting → A/R Aging')
    body(doc, 'Outstanding customer invoices grouped by how long they have been unpaid.')
    table(doc,
        ['Bucket', 'Meaning', 'Required Action'],
        [
            ['0–30 Days',  'Current',             'Monitor weekly'],
            ['31–60 Days', 'Slightly overdue',    'Send payment reminder via Sales Manager'],
            ['61–90 Days', 'Significantly overdue','Escalate to Sales Manager and MD'],
            ['90+ Days',       'Seriously overdue',   'Legal / MD escalation — immediate action'],
        ], col_widths=[1.3, 2.0, 3.2])

    pb(doc)
    h1(doc, '12.  Bank Reconciliation')
    body(doc, 'Path: Accounting → Bank Reconciliation  —  NEW in v2.0')
    body(doc, 'Match bank statement transactions against journal entries in the ERP to confirm the company\'s cash book agrees with the bank.')
    h2(doc, 'Process')
    bullet(doc, '1. Select the Bank Account from the dropdown.')
    bullet(doc, '2. Click Upload to import your bank statement (CSV format from your bank).')
    bullet(doc, '3. The system displays all transactions with a match status.')
    bullet(doc, '4. For Unmatched items, click Match to link to the corresponding journal entry.')
    table(doc,
        ['Match Status', 'Meaning', 'Action'],
        [
            ['Matched',   'Bank transaction confirmed against ERP',          'No action needed'],
            ['Unmatched', 'No corresponding ERP record found',               'Investigate immediately'],
            ['Partial',   'Amounts do not fully agree',                      'Review and correct'],
            ['Disputed',  'Transaction queried — under review',         'Escalate to management'],
        ], col_widths=[1.3, 2.7, 2.5])
    note(doc, 'Run bank reconciliation at least once per week. All Unmatched items must be resolved before month-end close.')

    h1(doc, '13.  Financial Statements')
    table(doc,
        ['Statement', 'Path'],
        [
            ['Comprehensive Income',  'Accounting → Comprehensive Income'],
            ['Financial Position',    'Accounting → Financial Position'],
            ['Equity Statement',      'Accounting → Equity'],
            ['Cash Flow Statement',   'Accounting → Cash Flow'],
            ['Accounting Notes',      'Accounting → Accounting Notes'],
        ], col_widths=[2.5, 4.0])

    pb(doc)
    h1(doc, '14.  Quick Reference')
    table(doc,
        ['Task', 'Path'],
        [
            ['Post a journal entry',        'Accounting → Journal Entries → + New Entry'],
            ['View general ledger',         'Accounting → General Ledger'],
            ['Run trial balance',           'Accounting → Trial Balance'],
            ['Record an expense',           'Accounting → Expenses → + New Expense'],
            ['Process payroll',             'Accounting → Payroll'],
            ['File a tax period',           'Accounting → Tax Periods'],
            ['Record petty cash',           'Accounting → Petty Cash → + New Entry'],
            ['Pay a supplier',              'Accounting → Supplier Payments → + New'],
            ['View A/R aging',              'Accounting → A/R Aging'],
            ['Reconcile bank statement',    'Accounting → Bank Reconciliation → Upload'],
            ['Generate financial reports',  'Accounting → Financial Reports'],
            ['View financial statements',   'Accounting → (Comprehensive Income / Financial Position / etc.)'],
            ['Send weekly report',          'A/R Aging page → Send Weekly Report'],
        ], col_widths=[2.8, 3.7])

    section_reminders(doc, [
        ('Bank reconciliation is mandatory weekly.', 'Unmatched transactions must be resolved before month-end close. Do not carry them forward.'),
        ('Invoices auto-post journals.', 'Do not manually re-post a journal for an already-issued invoice — this will create a duplicate entry.'),
        ('Confirmed receipts cannot be reversed by deletion.', 'Contact the Super Admin to void a receipt recorded in error.'),
        ('Month-end process order.', 'Reconcile bank → close petty cash → verify A/R aging → run trial balance → generate financial statements.'),
    ])
    footer(doc, 'Accounting Department')
    doc.save(os.path.join(OUT_DIR, '06_Accounting_Manual.docx'))
    print('Saved: 06_Accounting_Manual.docx')

# ════════════════════════════════════════════════════════════════════════════
# 7. HUMAN RESOURCES
# ════════════════════════════════════════════════════════════════════════════
def build_hr():
    doc = new_doc()
    cover_page(doc, 'Human Resources', 'For: HR Manager', 'db2777')
    section_getting_started(doc)

    h1(doc, '2.  Employees')
    body(doc, 'Path: HR → Employees')
    body(doc, 'The official record of all Aspee Pharmaceuticals staff. Every person employed by the company must have a record here.')
    h2(doc, 'To add an employee')
    bullet(doc, 'Click + Add Employee.')
    bullet(doc, 'Enter: Full Name, Employee ID, Department, Job Title, Date of Employment, Phone, and Emergency Contact.')
    bullet(doc, 'Save.')
    table(doc,
        ['Field', 'Example'],
        [
            ['Employee ID',       'EMP-2024-018'],
            ['Full Name',         'Abena Mensah'],
            ['Department',        'Quality Assurance'],
            ['Job Title',         'QA Analyst'],
            ['Date Employed',     '15 June 2024'],
            ['Phone',             '0244 987 654'],
            ['Emergency Contact', 'Mr. Kofi Mensah (Husband) — 0201 123 456'],
        ], col_widths=[2.3, 4.2])

    pb(doc)
    h1(doc, '3.  Attendance')
    body(doc, 'Path: HR → Attendance')
    body(doc, 'Daily attendance records for all staff.')
    table(doc,
        ['Field', 'Description'],
        [
            ['Employee',    'Select from registered employees list'],
            ['Date',        'The attendance date'],
            ['Check-In',    'Time of arrival'],
            ['Check-Out',   'Time of departure'],
            ['Hours Worked','Calculated automatically from check-in/out'],
            ['Status',      'Present / Absent / Late'],
        ], col_widths=[2.0, 4.5])

    h1(doc, '4.  Leave Management')
    body(doc, 'Path: HR → Leave Management')
    body(doc, 'Track and approve all staff leave requests. Approved leave updates the employee\'s leave balance automatically.')
    h2(doc, 'Leave types')
    table(doc,
        ['Leave Type', 'Notes'],
        [
            ['Annual Leave',             'Based on employee entitlement per company policy'],
            ['Sick Leave',               'Medical certificate required for extended periods'],
            ['Maternity / Paternity',    'As per Ghana Labour Act entitlements'],
            ['Compassionate Leave',      'Bereavement or family emergency'],
        ], col_widths=[2.3, 4.2])
    h2(doc, 'Approval process')
    bullet(doc, '1. Employee (or HR on their behalf) submits a leave request.')
    bullet(doc, '2. HR Manager reviews and approves or rejects.')
    bullet(doc, '3. Approved leave is deducted from the employee\'s leave balance.')

    h1(doc, '5.  Payroll Preparation')
    body(doc, 'Path: HR → Payroll Preparation')
    body(doc, 'HR prepares payroll input data — attendance records, leave deductions, and overtime — which is then processed by the Accounting department.')
    note(doc, 'Payroll preparation must be submitted to Accounting by the agreed monthly cut-off date. Late submissions delay salary payments.')

    pb(doc)
    h1(doc, '6.  Quick Reference')
    table(doc,
        ['Task', 'Path'],
        [
            ['Add a new employee',         'HR → Employees → + Add Employee'],
            ['Edit employee details',      'HR → Employees → Pencil icon on row'],
            ['Record attendance',          'HR → Attendance → + Add Record'],
            ['Submit leave request',       'HR → Leave Management → + New Request'],
            ['Approve / reject leave',     'HR → Leave Management → Review Request'],
            ['Prepare payroll data',       'HR → Payroll Preparation'],
            ['Send weekly HR report',      'HR → Employees page → Send Weekly Report'],
        ], col_widths=[2.8, 3.7])

    section_reminders(doc, [
        ('Employee IDs must be unique.', 'Once assigned, an Employee ID should never be reused — even for rehires.'),
        ('Payroll cut-off dates are firm.', 'Attend coordination meetings and submit HR data to Accounting on time each month.'),
        ('Leave balances are system-tracked.', 'Do not approve leave verbally without recording it. All approvals must be in the ERP so balances remain accurate.'),
        ('All actions are logged.', 'Every record creation, edit, and approval is timestamped with your name in the Audit Trail.'),
    ])
    footer(doc, 'Human Resources')
    doc.save(os.path.join(OUT_DIR, '07_HR_Manual.docx'))
    print('Saved: 07_HR_Manual.docx')

# ════════════════════════════════════════════════════════════════════════════
# 8. INTERNAL AUDIT
# ════════════════════════════════════════════════════════════════════════════
def build_internal_audit():
    doc = new_doc()
    cover_page(doc, 'Internal Audit', 'For: Internal Auditor', '6d28d9')
    section_getting_started(doc)

    h1(doc, '2.  Audit Plans')
    body(doc, 'Path: Internal Audit → Audit Plans')
    body(doc, 'Plan and schedule internal audit activities across all departments in advance.')
    h2(doc, 'Fields')
    table(doc,
        ['Field', 'Description', 'Example'],
        [
            ['Audit Title',   'Descriptive name for the audit',     'Q2 2026 Stores Inventory Audit'],
            ['Department',    'Department being audited',            'Stores'],
            ['Planned Date',  'Scheduled date for the audit',       '30 April 2026'],
            ['Auditor',       'Name of assigned auditor',           'Emmanuel Darko'],
            ['Scope',         'Objectives and boundaries',          'Reconcile physical count vs ERP stock records'],
        ], col_widths=[1.5, 2.0, 3.0])

    h1(doc, '3.  Audit Reports')
    body(doc, 'Path: Internal Audit → Audit Reports')
    body(doc, 'Record and store findings and recommendations from completed audits.')
    h2(doc, 'Report structure')
    table(doc,
        ['Section', 'Content'],
        [
            ['Findings',            'What was observed during the audit'],
            ['Risk Rating',         'High / Medium / Low — based on impact and likelihood'],
            ['Recommendations',     'Corrective actions proposed by the auditor'],
            ['Management Response', 'Department head\'s response and commitment'],
            ['Target Date',         'Agreed resolution date for each finding'],
        ], col_widths=[2.3, 4.2])

    pb(doc)
    h1(doc, '4.  Non-Conformances (NCs)')
    body(doc, 'Path: Internal Audit → Non-Conformances')
    body(doc, 'A Non-Conformance is any deviation from established procedures, standards, or regulatory requirements.')
    table(doc,
        ['Severity', 'Meaning', 'Response Required'],
        [
            ['Minor',    'Isolated, low-risk deviation',                             'Document and monitor'],
            ['Major',    'Systematic or high-risk deviation',                        'Root-cause analysis + corrective action plan'],
            ['Critical', 'Immediate risk to product quality, patient safety, or regulatory compliance',
                         'Stop affected operations + immediate management notification'],
        ], col_widths=[1.0, 2.7, 2.8])
    warning(doc, 'Critical Non-Conformances must be escalated to the QA Manager and Managing Director immediately. Do not wait for the next scheduled report.')

    h1(doc, '5.  Audit Trail (Read Access)')
    body(doc, 'Path: Settings → Audit Trail')
    body(doc, 'The Internal Auditor has read access to the system-wide Audit Trail — a complete log of every create, edit, delete, and approve action performed by any user.')
    note(doc, 'Use the Audit Trail to verify whether an action was actually performed, when, and by whom. This is your primary evidence source for audit findings.')

    h1(doc, '6.  Quick Reference')
    table(doc,
        ['Task', 'Path'],
        [
            ['Create an audit plan',         'Internal Audit → Audit Plans → + New Plan'],
            ['Record audit findings',        'Internal Audit → Audit Reports → + New Report'],
            ['Log a non-conformance',        'Internal Audit → Non-Conformances → + New NC'],
            ['View system audit trail',      'Settings → Audit Trail'],
            ['Send weekly report',           'Internal Audit → Audit Plans page → Send Weekly Report'],
        ], col_widths=[2.8, 3.7])

    section_reminders(doc, [
        ('Independence is critical.', 'Do not audit a process you participated in. Maintain objectivity at all times.'),
        ('Evidence first.', 'Every audit finding must be supported by evidence from the system (Audit Trail records, screenshots, data exports).'),
        ('Critical NCs require immediate escalation.', 'Do not wait for report cycles — notify QA and the Managing Director the same day.'),
        ('Close loops.', 'Follow up on all open findings until they are resolved and evidence of closure is documented.'),
    ])
    footer(doc, 'Internal Audit')
    doc.save(os.path.join(OUT_DIR, '08_Internal_Audit_Manual.docx'))
    print('Saved: 08_Internal_Audit_Manual.docx')

# ════════════════════════════════════════════════════════════════════════════
# 9. COMPLIANCE & REGULATORS
# ════════════════════════════════════════════════════════════════════════════
def build_compliance():
    doc = new_doc()
    cover_page(doc, 'Compliance & Regulators', 'For: Quality Assurance Officer / Managing Director', '059669')
    section_getting_started(doc)

    h1(doc, '2.  Regulators & Renewals')
    body(doc, 'Path: Compliance → Regulators & Renewals')
    body(doc, 'Track all regulatory licences, certifications, and permit renewals to ensure Aspee Pharmaceuticals remains fully compliant with Ghana FDA and all other relevant bodies.')
    h2(doc, 'Fields to maintain')
    table(doc,
        ['Field', 'Description', 'Example'],
        [
            ['Regulator',      'Name of the regulatory authority',      'FDA Ghana'],
            ['Licence Type',   'Type of permit or certification',       "Manufacturer's Licence"],
            ['Licence Number', 'Official licence/permit reference',     'MFR-GH-2024-0089'],
            ['Issue Date',     'Date the licence was granted',          '01 Jan 2024'],
            ['Expiry Date',    'Date the licence expires',              '31 Dec 2026'],
            ['Renewal Status', 'Current renewal state',                 'Active / Due for Renewal / Expired'],
        ], col_widths=[1.8, 2.2, 2.5])
    h3(doc, 'Current Licence Register')
    table(doc,
        ['Regulator', 'Licence Type', 'Number', 'Expiry', 'Status'],
        [
            ['FDA Ghana', "Manufacturer's Licence", 'MFR-GH-2024-0089', '31 Dec 2026', 'Active'],
            ['EPA',       'Factory Registration',   'EPA-2023-FAC-1102','30 Jun 2026', 'Due for Renewal'],
            ['GSS',       'Business Registration',  'GSS-REG-20190',    'Permanent',   'Active'],
        ], col_widths=[1.4, 1.8, 1.7, 1.3, 1.3])

    pb(doc)
    h1(doc, '3.  Renewal Management Process')
    h2(doc, 'Recommended renewal timeline')
    table(doc,
        ['Days Before Expiry', 'Action'],
        [
            ['90 days', 'Initiate contact with the regulatory authority. Begin assembling required documents.'],
            ['60 days', 'Submit renewal application with all supporting documents.'],
            ['30 days', 'Follow up on application status. Escalate if not acknowledged.'],
            ['0 days',  'Licence expires. Operations may be suspended. Immediate escalation to MD required.'],
        ], col_widths=[1.5, 5.0])
    warning(doc, 'An expired Manufacturer\'s Licence from FDA Ghana means Aspee Pharmaceuticals cannot legally manufacture or sell pharmaceutical products. This is a critical risk — never let this lapse.')

    h1(doc, '4.  Non-Conformances Related to Compliance')
    body(doc, 'If a regulatory inspection finds a compliance issue, log a Non-Conformance in the Internal Audit module immediately.')
    bullet(doc, 'Path: Internal Audit → Non-Conformances → + New NC')
    bullet(doc, 'Set severity to Major or Critical as appropriate.')
    bullet(doc, 'Assign a corrective action with a target resolution date.')
    bullet(doc, 'Update the NC record as each corrective action is completed.')

    h1(doc, '5.  Quick Reference')
    table(doc,
        ['Task', 'Path'],
        [
            ['View all licences',         'Compliance → Regulators & Renewals'],
            ['Add a new licence record',  'Compliance → Regulators & Renewals → + New'],
            ['Update renewal status',     'Compliance → Regulators & Renewals → Edit row'],
            ['Log a compliance NC',       'Internal Audit → Non-Conformances → + New NC'],
        ], col_widths=[2.8, 3.7])

    section_reminders(doc, [
        ('90-day rule.', 'Start every renewal at least 90 days before expiry. Regulatory bodies in Ghana can have slow processing times.'),
        ('Expired licences must be escalated to the MD immediately.', 'Do not attempt to continue operations under an expired licence.'),
        ('Keep supporting documents attached.', 'Attach the actual licence document to each record so it is accessible during inspections.'),
        ('All changes are logged.', 'Every update to the compliance register is recorded in the Audit Trail.'),
    ])
    footer(doc, 'Compliance & Regulators')
    doc.save(os.path.join(OUT_DIR, '09_Compliance_Manual.docx'))
    print('Saved: 09_Compliance_Manual.docx')

# ════════════════════════════════════════════════════════════════════════════
# 10. SUPER ADMIN / SETTINGS
# ════════════════════════════════════════════════════════════════════════════
def build_admin():
    doc = new_doc()
    cover_page(doc, 'System Administration', 'For: Super Admin / IT Administrator', '1e3a5f')
    section_getting_started(doc)

    h1(doc, '2.  User Management')
    body(doc, 'Path: Settings → User Management')
    body(doc, 'Create system users, assign roles, and deactivate accounts. Only Super Admins can access this page.')
    h2(doc, 'To create a new user')
    bullet(doc, 'Click + Add User.')
    bullet(doc, 'Enter the staff member\'s company email address.')
    bullet(doc, 'Select their Role from the dropdown.')
    bullet(doc, 'Save.')
    note(doc, 'The user must also be registered in Supabase Auth (via Auth dashboard or invitation email) before they can log in. This is a separate step handled by IT.')
    h2(doc, 'Available roles and permissions')
    table(doc,
        ['Role', 'Accessible Modules'],
        [
            ['Super Admin',        'Everything — full system access'],
            ['Managing Director',  'Dashboard, Weekly Reports Review'],
            ['Sales Manager',      'Sales (all), Customers, Weekly Reports'],
            ['Van Sales Rep',      'Invoices, Dispatch, Receipts, Weekly Reports'],
            ['Purchasing Manager', 'Procurement, Supplier Payments, GRN, Weekly Reports'],
            ['Store Manager',      'Stores (all), Production, Weekly Reports'],
            ['Production Manager', 'Production, Weekly Reports'],
            ['Quality Assurance',  'QA (all), GRN, Compliance, Weekly Reports'],
            ['Accountant',         'Accounting (all), Collections, Supplier Payments, Weekly Reports'],
            ['HR Manager',         'HR (all), Weekly Reports'],
            ['Internal Auditor',   'Internal Audit, Weekly Reports'],
        ], col_widths=[2.0, 4.5])

    pb(doc)
    h1(doc, '3.  Profile Settings')
    body(doc, 'Path: Settings → Profile')
    body(doc, 'Available to all users. Staff can update their display name, contact details, and profile photo.')

    h1(doc, '4.  Report Settings')
    body(doc, 'Path: Settings → Report Settings')
    body(doc, 'Configure the company details that appear on all printed documents:')
    bullet(doc, 'Company name and registration number')
    bullet(doc, 'Company address and contact details')
    bullet(doc, 'Company logo (uploaded as an image)')
    bullet(doc, 'Footer text for invoices and reports')

    h1(doc, '5.  Audit Trail')
    body(doc, 'Path: Settings → Audit Trail')
    body(doc, 'A complete, immutable log of every action performed in the system by any user.')
    table(doc,
        ['Field', 'Description'],
        [
            ['User',       'The staff member who performed the action'],
            ['Action',     'Create / Update / Delete / Approve'],
            ['Module',     'Which part of the system was affected'],
            ['Record',     'The specific record that was changed'],
            ['Timestamp',  'Exact date and time of the action'],
        ], col_widths=[1.5, 5.0])
    warning(doc, 'The Audit Trail cannot be edited or deleted by any user, including Super Admins. It is the definitive accountability record.')

    pb(doc)
    h1(doc, '6.  System-Wide Reminders for Admins')
    h2(doc, 'Onboarding a new user — checklist')
    bullet(doc, '[ ] Create auth account in Supabase (invite email or manual creation).')
    bullet(doc, '[ ] Add user record in Settings → User Management with correct role.')
    bullet(doc, '[ ] Confirm the user can log in and see only their permitted modules.')
    bullet(doc, '[ ] Add the user to any department-specific notification groups if applicable.')
    h2(doc, 'Offboarding a user — checklist')
    bullet(doc, '[ ] Disable the auth account in Supabase immediately on departure.')
    bullet(doc, '[ ] Reassign any open tasks or pending approvals to another user.')
    bullet(doc, '[ ] Archive the user record in User Management.')

    h1(doc, '7.  Quick Reference')
    table(doc,
        ['Task', 'Path'],
        [
            ['Create a new user',          'Settings → User Management → + Add User'],
            ['Change a user\'s role',      'Settings → User Management → Edit row'],
            ['Update company logo/details','Settings → Report Settings'],
            ['View full audit trail',      'Settings → Audit Trail'],
            ['Reset a user password',      'Supabase Auth dashboard (outside ERP)'],
            ['View any module',            'Super Admin has access to all sidebar sections'],
        ], col_widths=[2.8, 3.7])

    section_reminders(doc, [
        ('Only assign roles that are needed.', 'The principle of least privilege applies. Do not assign Super Admin unless the person is an IT administrator.'),
        ('Disable accounts immediately on staff departure.', 'An active account belonging to a former employee is a security risk.'),
        ('Never share the Super Admin credentials.', 'Each administrator must have their own account — shared credentials break the Audit Trail.'),
        ('Test after every role assignment.', 'Log in as the new user (or use a test account with that role) to verify they see only what they should.'),
    ])
    footer(doc, 'System Administration')
    doc.save(os.path.join(OUT_DIR, '10_Admin_Manual.docx'))
    print('Saved: 10_Admin_Manual.docx')

# ════════════════════════════════════════════════════════════════════════════
# RUN ALL
# ════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    build_procurement()
    build_qa()
    build_stores()
    build_production()
    build_sales()
    build_accounting()
    build_hr()
    build_internal_audit()
    build_compliance()
    build_admin()
    print(f'\nAll 10 department manuals saved to:\n{OUT_DIR}')
